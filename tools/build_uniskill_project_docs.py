from __future__ import annotations

from collections import OrderedDict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\MOHIT\Desktop\uniskill")
DOWNLOADS = Path(r"C:\Users\MOHIT\Downloads")
GENERATED = ROOT / "generated_docs"
FINAL_DIR = GENERATED / "final"
ASSETS = GENERATED / "assets"
SCREENSHOTS = GENERATED / "screenshots"

SOURCE_SYNOPSIS = GENERATED / "SYNOPSIS_FINAL_uniskill_cleaned.docx"
SOURCE_SRS = GENERATED / "SRS_FINAL_uniskill_cleaned.docx"
SOURCE_UML = GENERATED / "UNISKILL_UML_FINAL.docx"

TEMPLATE_SYNOPSIS = DOWNLOADS / "Synopsis Format.docx"
TEMPLATE_SRS = DOWNLOADS / "SRS_template 2026.docx"
TEMPLATE_REPORT = DOWNLOADS / "BTech_CSE_MINI Project- Report Format.docx"
TEMPLATE_SDS = DOWNLOADS / "(UML Diag)Software Design Specification Template.docx"

PROJECT_TITLE = "UNISKILL"
PROJECT_LONG_TITLE = "UniSkill: Verified Student Skill Exchange Platform"
MONTH_YEAR = "APRIL 2026"
STUDENT_PLACEHOLDER = "Name(s) of Student(s) to be inserted by submitting team"
ROLL_PLACEHOLDER = "University Roll No(s) to be inserted by submitting team"
GUIDE_PLACEHOLDER = "Guide Name(s) to be inserted"


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def is_heading(paragraph) -> bool:
    style = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()
    return bool(text) and (
        style.startswith("Heading")
        or style in {"TOCEntry", "requirement", "level 4", "level 3 text", "toc 1"}
    )


def find_paragraphs(doc: Document, text: str) -> list:
    return [p for p in doc.paragraphs if p.text.strip() == text]


def find_first(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"Paragraph not found: {text}")


def find_index(doc: Document, text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text:
            return i
    raise ValueError(f"Paragraph index not found: {text}")


def replace_all(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)


def insert_paragraph_before(anchor, text: str = "", style: str | None = None, bold: bool = False, italic: bool = False):
    p = anchor.insert_paragraph_before("")
    if style:
        try:
            p.style = style
        except KeyError:
            pass
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def insert_image_before(anchor, image_path: Path, width: float = 6.4, caption: str | None = None):
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    if caption:
        cp = anchor.insert_paragraph_before(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cp.runs:
            cp.runs[0].italic = True


def insert_table_before(doc: Document, anchor, headers: list[str], rows: list[list[str]], style: str = "Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = style
    except KeyError:
        pass
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    anchor._p.addprevious(table._tbl)
    spacer = anchor.insert_paragraph_before("")
    return table


def clear_between(doc: Document, start_text: str, end_text: str) -> None:
    start = find_first(doc, start_text)
    end = find_first(doc, end_text)
    started = False
    for p in list(doc.paragraphs):
        if p == start:
            started = True
            continue
        if p == end:
            break
        if started:
            delete_paragraph(p)


def remove_instructional_paragraphs(doc: Document, extra_contains: list[str] | None = None) -> None:
    extra_contains = extra_contains or []
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if text.startswith("<"):
            delete_paragraph(p)
            continue
        if any(marker in text for marker in extra_contains):
            delete_paragraph(p)


def trim_after_first_prefix(doc: Document, prefix: str) -> None:
    found = False
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if found:
            delete_paragraph(p)
            continue
        if text.startswith(prefix):
            found = True


def trim_body_after_paragraph_prefix(doc: Document, prefix: str) -> None:
    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            target = p._p
            break
    if target is None:
        return
    body = target.getparent()
    seen = False
    for child in list(body):
        if child is target:
            seen = True
            continue
        if seen and child.tag.endswith("sectPr"):
            continue
        if seen:
            body.remove(child)


def insert_blocks_before(doc: Document, end_anchor, blocks: list[tuple]):
    for block in reversed(blocks):
        kind = block[0]
        if kind == "p":
            _, text, style = block
            insert_paragraph_before(end_anchor, text, style)
        elif kind == "bullet":
            _, items = block
            for item in reversed(items):
                insert_paragraph_before(end_anchor, item, "List Paragraph")
        elif kind == "img":
            _, path, width, caption = block
            insert_image_before(end_anchor, path, width, caption)
        elif kind == "table":
            _, headers, rows = block
            insert_table_before(doc, end_anchor, headers, rows)


def extract_sections_by_h1(path: Path) -> OrderedDict[str, list[str]]:
    doc = Document(path)
    sections: OrderedDict[str, list[str]] = OrderedDict()
    current = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            current = text
            sections[current] = []
        elif current and not style.startswith("Heading"):
            sections[current].append(text)
    return sections


def extract_section_texts(path: Path, heading_text: str) -> list[str]:
    doc = Document(path)
    capture = False
    out: list[str] = []
    start_style_level = None
    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if text == heading_text and style.startswith("Heading"):
            capture = True
            start_style_level = int(style.split()[-1])
            continue
        if capture and style.startswith("Heading") and text:
            try:
                level = int(style.split()[-1])
            except Exception:
                level = 9
            if level <= start_style_level:
                break
        if capture and text and not style.startswith("Heading"):
            out.append(text)
    return out


def compact(texts: list[str], limit: int | None = None) -> list[str]:
    cleaned = [t for t in texts if t and not t.startswith("UniSkill approaches ")]
    return cleaned[:limit] if limit else cleaned


def build_synopsis() -> Path:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE_SYNOPSIS)
    replace_all(doc, "Synopsis Template B.Tech (VII sem) Project Work-1", "")
    replace_all(doc, "<TITLE OF THE PROJECT>", PROJECT_LONG_TITLE.upper())
    replace_all(doc, "<Name of the Student(s)>", STUDENT_PLACEHOLDER)
    replace_all(doc, "<University Roll No(s) of the Student(s)>", ROLL_PLACEHOLDER)
    replace_all(doc, "<Name of the Guide(s)>", GUIDE_PLACEHOLDER)
    replace_all(doc, "<MONTH YEAR>", MONTH_YEAR)

    # Remove notes/examples from the template body.
    remove = False
    for p in list(doc.paragraphs):
        if p.text.strip() == "Note:":
            remove = True
        if remove:
            delete_paragraph(p)

    source = extract_sections_by_h1(SOURCE_SYNOPSIS)

    doc.add_page_break()
    sections = [
        ("Introduction", compact(source["1. Introduction"] + source["3. Problem Definition"][:4])),
        ("Literature Review", compact(source["2. Literature Review"] + source["Appendix A: Stakeholder, Market, and Rollout Analysis"][:12])),
        ("Problem Definition", compact(source["3. Problem Definition"] + source["2. Literature Review"][-3:] + source["Appendix A: Stakeholder, Market, and Rollout Analysis"][12:32])),
        ("Objectives", compact(source["4. Objectives"])),
        (
            "Methodology",
            compact(
                source["5. Scope and Feasibility"]
                + source["6. Existing and Proposed System"]
                + source["7. Methodology and Architecture"]
                + source["8. Module Description"]
                + source["9. Requirements and Tools"]
                + source["10. Risks, Testing, and Deployment Strategy"]
                + source["11. Expected Outcomes and Future Enhancement"]
                + source["Appendix A: Stakeholder, Market, and Rollout Analysis"][32:40]
                + source["Appendix B: Detailed Module Walkthroughs and Scenario Narratives"][:40]
                + source["Appendix D: Testing Scenarios and Expected Observations"][:40]
                + source["Appendix C: Demo Flow and Evaluation Readiness"]
            ),
        ),
    ]

    for title, paragraphs in sections:
        p = doc.add_paragraph(title)
        p.style = "Heading 1"
        for text in paragraphs:
            doc.add_paragraph(text)
        if title == "Methodology":
            method_table_headers = ["Layer", "Technology", "Role in UniSkill"]
            method_table_rows = [
                ["Frontend", "Next.js 16, React 19, Tailwind CSS 4", "Public website, onboarding, profile and marketplace experiences"],
                ["Backend", "Express 5, TypeScript, Socket.io", "OTP, wallet, sessions, realtime signaling and business routes"],
                ["Platform Services", "Supabase Auth, PostgreSQL, RLS, Storage", "Identity, persistence, protected data access and storage"],
            ]
            table = doc.add_table(rows=1, cols=3)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            for i, h in enumerate(method_table_headers):
                table.rows[0].cells[i].text = h
            for row in method_table_rows:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = value
            doc.add_paragraph("")
            for image_path, caption in [
                (ASSETS / "architecture_overview.png", "Figure: UniSkill logical architecture overview"),
                (SCREENSHOTS / "home.png", "Figure: UniSkill landing page"),
                (ASSETS / "use_case.png", "Figure: UniSkill use case overview"),
                (ASSETS / "dfd.png", "Figure: UniSkill data flow view"),
                (SCREENSHOTS / "login.png", "Figure: UniSkill login page"),
                (SCREENSHOTS / "signup.png", "Figure: UniSkill signup page"),
            ]:
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(str(image_path), width=Inches(6.3))
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cap.runs:
                    cap.runs[0].italic = True

    p = doc.add_paragraph("References")
    p.style = "Heading 1"
    refs = [
        "[1] Vercel, Next.js Documentation, https://nextjs.org/docs, accessed April 28, 2026.",
        "[2] React Documentation, https://react.dev, accessed April 28, 2026.",
        "[3] Supabase Documentation, https://supabase.com/docs, accessed April 28, 2026.",
        "[4] Express.js Documentation, https://expressjs.com, accessed April 28, 2026.",
        "[5] Prisma Documentation, https://docs.prisma.io, accessed April 28, 2026.",
        "[6] Socket.IO Documentation, https://socket.io, accessed April 28, 2026.",
        "[7] Internal project references: frontend/package.json, backend/package.json, supabase migrations, and generated UniSkill design assets.",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    out = FINAL_DIR / "UNISKILL_Synopsis_Template_Filled.docx"
    doc.save(out)
    return out


def build_srs() -> Path:
    doc = Document(TEMPLATE_SRS)
    srs_sections = extract_sections_by_h1(SOURCE_SRS)

    section_map = {
        "Purpose": compact(extract_section_texts(SOURCE_SRS, "1.1 Purpose"), 2),
        "Document Conventions": compact(extract_section_texts(SOURCE_SRS, "1.2 Document Conventions"), 2),
        "Intended Audience and Reading Suggestions": compact(extract_section_texts(SOURCE_SRS, "1.3 Intended Audience and Reading Suggestions"), 2),
        "Product Scope": compact(extract_section_texts(SOURCE_SRS, "1.4 Product Scope"), 2),
        "Product Perspective": compact(extract_section_texts(SOURCE_SRS, "2.1 Product Perspective"), 3),
        "Product Functions": compact(extract_section_texts(SOURCE_SRS, "2.2 Product Functions"), 4),
        "User Classes and Characteristics": compact(extract_section_texts(SOURCE_SRS, "2.3 User Classes and Characteristics"), 4),
        "Operating Environment": compact(extract_section_texts(SOURCE_SRS, "2.4 Operating Environment"), 4),
        "Design and Implementation Constraints": compact(extract_section_texts(SOURCE_SRS, "2.5 Design and Implementation Constraints"), 4),
        "User Documentation": compact(extract_section_texts(SOURCE_SRS, "2.6 User Documentation"), 3),
        "Assumptions and Dependencies": compact(extract_section_texts(SOURCE_SRS, "2.7 Assumptions and Dependencies"), 4),
        "User Interfaces": compact(extract_section_texts(SOURCE_SRS, "3.1 User Interfaces") + srs_sections["Appendix A: Detailed Use Case Specifications"][:8], 13),
        "Hardware Interfaces": compact(extract_section_texts(SOURCE_SRS, "3.2 Hardware Interfaces"), 2),
        "Software Interfaces": compact(extract_section_texts(SOURCE_SRS, "3.3 Software Interfaces") + srs_sections["Appendix D: Detailed Data Rules and Quality Constraints"][:10], 15),
        "Communications Interfaces": compact(extract_section_texts(SOURCE_SRS, "3.4 Communications Interfaces"), 3),
        "Performance Requirements": compact(extract_section_texts(SOURCE_SRS, "5.1 Performance") + srs_sections["Appendix C: Error Handling, Boundary Conditions, and Recovery Requirements"][8:16], 12),
        "Safety Requirements": [
            "UniSkill does not control hazardous physical equipment, but it must prevent user harm caused by identity abuse, double charging, unauthorized room access, or misleading service claims.",
            "The platform shall preserve safe defaults by requiring authentication for protected surfaces, logging sensitive mutations, and preventing duplicate credit movement during session settlement.",
        ],
        "Security Requirements": compact(extract_section_texts(SOURCE_SRS, "5.2 Security") + srs_sections["Appendix C: Error Handling, Boundary Conditions, and Recovery Requirements"][:8], 13),
        "Software Quality Attributes": compact(
            extract_section_texts(SOURCE_SRS, "5.3 Reliability")
            + extract_section_texts(SOURCE_SRS, "5.4 Usability")
            + extract_section_texts(SOURCE_SRS, "5.5 Scalability")
            + extract_section_texts(SOURCE_SRS, "5.6 Maintainability"),
            8,
        ),
    }
    section_map["Software Quality Attributes"] = compact(
        section_map["Software Quality Attributes"] + srs_sections["Appendix D: Detailed Data Rules and Quality Constraints"][:12],
        20,
    )

    heading_order = [
        "Purpose",
        "Document Conventions",
        "Intended Audience and Reading Suggestions",
        "Product Scope",
        "Product Perspective",
        "Product Functions",
        "User Classes and Characteristics",
        "Operating Environment",
        "Design and Implementation Constraints",
        "User Documentation",
        "Assumptions and Dependencies",
        "User Interfaces",
        "Hardware Interfaces",
        "Software Interfaces",
        "Communications Interfaces",
        "Performance Requirements",
        "Safety Requirements",
        "Security Requirements",
        "Software Quality Attributes",
    ]

    for heading in heading_order:
        start = find_first(doc, heading)
        idx = find_index(doc, heading)
        end = None
        for candidate in doc.paragraphs[idx + 1 :]:
            if is_heading(candidate):
                end = candidate
                break
        if end is None:
            continue
        clear_between(doc, heading, end.text.strip())
        blocks: list[tuple] = []
        for para in section_map[heading]:
            blocks.append(("p", para, "Normal"))
        if heading == "Product Perspective":
            blocks.append(("img", ASSETS / "architecture_overview.png", 6.2, "Figure: UniSkill system context"))
            blocks.append(("img", ASSETS / "component_design.png", 6.0, "Figure: UniSkill component interaction overview"))
        if heading == "Product Functions":
            blocks.append(
                (
                    "table",
                    ["Function Group", "Primary Output", "Related Surface"],
                    [
                        ["Registration and verification", "Trusted campus account", "Signup, OTP, verification vault"],
                        ["Marketplace discovery", "Searchable mentor and service cards", "Explore, profile, service listing APIs"],
                        ["Session lifecycle", "Booked, confirmed and completed sessions", "Sessions API, room page, dashboard"],
                        ["Wallet and credits", "Audited balance changes", "Wallet page, transactions table, server-side checks"],
                    ],
                )
            )
        if heading == "User Interfaces":
            blocks.append(("img", SCREENSHOTS / "home.png", 5.9, "Figure: Public interface of UniSkill"))
            blocks.append(("img", SCREENSHOTS / "login.png", 5.9, "Figure: Login interface of UniSkill"))
            blocks.append(("img", SCREENSHOTS / "signup.png", 5.9, "Figure: Signup interface of UniSkill"))
        if heading == "Software Interfaces":
            blocks.append(("img", ASSETS / "erd.png", 6.0, "Figure: Data-facing software interface context"))
            blocks.append(("img", ASSETS / "dfd.png", 6.0, "Figure: Functional data flow between interfaces"))
        if heading == "Communications Interfaces":
            blocks.append(("img", ASSETS / "booking_sequence.png", 6.0, "Figure: Communication sequence around booking and room coordination"))
        if heading == "Security Requirements":
            blocks.append(("img", ASSETS / "state_transition.png", 6.0, "Figure: Session state control relevant to secure transitions"))
        insert_blocks_before(doc, end, blocks)

    # Replace features section.
    clear_between(doc, "System Features", "Nonfunctional Requirements")
    end_anchor = find_first(doc, "Nonfunctional Requirements")
    feature_blocks: list[tuple] = []
    features = [
        (
            "User Registration, Verification, and Profile Onboarding",
            [
                "UniSkill shall allow a student to create an account, verify identity signals, and establish a profile that can later be used for teaching or learning. Priority: High.",
                "Typical stimulus/response flow: user enters identity data, system validates format and uniqueness, OTP routes verify the supplied contact method, and the profile plus vault state are persisted for future use.",
                "REQ-1: The system shall support account creation using student-facing registration flows.",
                "REQ-2: The system shall support OTP-based verification for protected onboarding steps.",
                "REQ-3: The system shall separate public profile data from sensitive verification-vault data.",
                "REQ-4: The system shall maintain verification step progress so onboarding can be resumed.",
            ],
        ),
        (
            "Marketplace Discovery and Service Listing Management",
            [
                "UniSkill shall allow mentors to publish teachable offerings and learners to discover those offerings through searchable marketplace views. Priority: High.",
                "Typical stimulus/response flow: mentor updates profile or listing, system saves structured metadata, learner filters or searches the marketplace, and the frontend renders suitable mentor or listing cards.",
                "REQ-5: The system shall allow verified users to create and maintain service listings.",
                "REQ-6: The system shall expose category, skill, price, and availability metadata for discovery.",
                "REQ-7: The system shall support search and filter operations for marketplace exploration.",
            ],
        ),
        (
            "Session Booking and Realtime Room Support",
            [
                "UniSkill shall transform a marketplace selection into a trackable session lifecycle, including booking, confirmation, room join readiness, and completion. Priority: High.",
                "Typical stimulus/response flow: learner books a session, sessions API creates a pending record, mentor receives notification, both users later access room interaction surfaces, and status updates remain traceable.",
                "REQ-8: The system shall create a pending session record for each valid booking request.",
                "REQ-9: The system shall maintain explicit session states such as pending, confirmed, in_progress, completed, cancelled, and disputed.",
                "REQ-10: The system shall support realtime communication or signaling hooks for live-room participation.",
            ],
        ),
        (
            "Wallet, Credits, and Transaction Audit",
            [
                "UniSkill shall control value exchange through a wallet and transaction model rather than relying on unaudited client-side changes. Priority: High.",
                "Typical stimulus/response flow: booking or transfer request reaches the server, wallet balance and session context are validated, a transaction row is persisted, and the frontend reflects the latest balance state.",
                "REQ-11: The system shall maintain a wallet balance for each active user profile.",
                "REQ-12: The system shall record transaction rows for credit, debit, refund, or bonus events.",
                "REQ-13: The system shall prevent duplicate or arbitrary balance mutation through server-side validation.",
            ],
        ),
        (
            "Reviews, Notifications, and Governance Support",
            [
                "UniSkill shall support trust-building and accountability through reviews, notifications, audit traces, and moderation-oriented data structures. Priority: Medium to High.",
                "Typical stimulus/response flow: completed interaction enables review submission, notification rows are created for important lifecycle changes, and audit or dispute-related records remain available for administrative follow-up.",
                "REQ-14: The system shall allow users to review completed learning interactions.",
                "REQ-15: The system shall maintain notification records for important product events.",
                "REQ-16: The system shall preserve auditability for sensitive changes and governance-related actions.",
            ],
        ),
    ]
    for title, lines in reversed(features):
        insert_paragraph_before(end_anchor, title, "Heading 2")
        insert_paragraph_before(end_anchor, "4.x.1 Description and Priority", "level 4")
        insert_paragraph_before(end_anchor, lines[0], "level 3 text")
        insert_paragraph_before(end_anchor, "4.x.2 Stimulus/Response Sequences", "level 4")
        insert_paragraph_before(end_anchor, lines[1], "level 3 text")
        insert_paragraph_before(end_anchor, "4.x.3 Functional Requirements", "level 4")
        for req in reversed(lines[2:]):
            insert_paragraph_before(end_anchor, req, "requirement")

    # Project plan and appendices.
    clear_between(doc, "Team Members", "Division of Work")
    insert_blocks_before(
        doc,
        find_first(doc, "Division of Work"),
        [("p", "Submitting team details are to be inserted before final academic submission. The documentation has been prepared to support one or more student authors working on the UniSkill platform.", "Normal")],
    )
    clear_between(doc, "Division of Work", "Time Schedule")
    insert_blocks_before(
        doc,
        find_first(doc, "Time Schedule"),
        [
            ("p", "Suggested division of work: requirements analysis and project report preparation; frontend experience design; backend route and data-layer development; Supabase schema, policy, and testing support. The exact member allocation should match the final student team.", "Normal")
        ],
    )
    clear_between(doc, "Time Schedule", "Appendix A: Glossary")
    insert_blocks_before(
        doc,
        find_first(doc, "Appendix A: Glossary"),
        [("bullet", ["Requirements analysis and planning: 2 weeks", "Database and architecture design: 2 weeks", "Frontend and backend implementation: 6 weeks", "Integration, testing, and documentation: 4 weeks"])],
    )
    clear_between(doc, "Appendix A: Glossary", "Appendix B: To Be Determined List")
    insert_blocks_before(
        doc,
        find_first(doc, "Appendix B: To Be Determined List"),
        [
            ("bullet", [
                "RLS: Row Level Security used to restrict database access by row ownership and role.",
                "OTP: One-Time Password used for verification-driven onboarding steps.",
                "Service listing: Structured mentor offer containing title, category, duration, and credit price.",
                "Verification vault: Private storage area for sensitive campus identity details.",
                "Realtime room: Session interaction surface for chat or live collaboration.",
            ])
        ],
    )
    # Replace Appendix B template paragraph with explicit TBD list.
    for p in list(doc.paragraphs):
        if "Collect a numbered list of the TBD" in p.text:
            p.text = "TBD-1: Final student names, roll numbers, and guide names are to be inserted before submission.\nTBD-2: Final production deployment host and environment-specific secrets.\nTBD-3: Whether an extended admin console will be delivered in the current academic milestone."

    remove_instructional_paragraphs(
        doc,
        [
            "This template illustrates organizing the functional requirements",
            "Don’t really say",
            "Don't really say",
            "Each requirement should be uniquely identified",
        ],
    )

    out = FINAL_DIR / "UNISKILL_SRS_Template_Filled.docx"
    doc.save(out)
    return out


def build_sds() -> Path:
    doc = Document(TEMPLATE_SDS)
    replace_all(doc, "Software Design Specification for <Project> based on event driven (function oriented) technology", f"Software Design Specification for {PROJECT_TITLE} based on event driven (function oriented) technology")

    # Remove the object-oriented template half.
    occurrences = [p for p in doc.paragraphs if p.text.strip() == "Software Design Specification Template"]
    if len(occurrences) > 1:
        second = occurrences[1]
        deleting = False
        for p in list(doc.paragraphs):
            if p == second:
                deleting = True
            if deleting:
                delete_paragraph(p)

    section_blocks = {
        "1.1 Purpose": [("p", "This SDS documents the event-driven and function-oriented design of UniSkill, a student skill exchange platform that combines onboarding, marketplace discovery, service listing, session booking, credit transfer, and realtime room support inside one campus-oriented workflow.", "Normal")],
        "1.2 Document Conventions": [("p", "This document uses numbered analysis and design sections aligned with the supplied SDS template. Figures are labeled in the order they appear, tables summarize structured entities or rules, and narrative descriptions focus on observable implementation behavior from the current UniSkill repository.", "Normal")],
        "1.3 Intended Audience and Reading Suggestions": [("p", "The SDS is intended for faculty evaluators, student developers, testers, and future maintainers. Readers should begin with the introduction, review the analysis diagrams to understand system behavior, and then move into the design model for implementation-level structure.", "Normal")],
        "1.4 References": [("bullet", [
            "Next.js Documentation: https://nextjs.org/docs",
            "React Documentation: https://react.dev",
            "Supabase Documentation: https://supabase.com/docs",
            "Express.js Documentation: https://expressjs.com",
            "Internal references: UniSkill frontend, backend, and Supabase migration files",
        ])],
        "2.1 Methodology used": [("p", "UniSkill follows a function-oriented, event-driven design approach. Important user events such as signup, OTP verification, marketplace search, booking, wallet settlement, review submission, and room entry trigger discrete backend flows that update persistent state and communicate changes back to the UI.", "Normal"), ("p", "Development is iterative: the frontend presents the interaction surface, app routes or backend routes validate requests, and Supabase-backed tables and policies preserve consistency. This makes the platform suitable for academic explanation because each business event can be traced to a concrete interface, route, and storage effect.", "Normal")],
        "2.2 Use Case diagram": [("p", "The figure below summarizes the primary actors and their major use cases in UniSkill.", "Normal"), ("img", ASSETS / "use_case.png", 6.3, "Figure: UniSkill use case view"), ("p", "Representative use case specification: Book Session. Precondition: learner is authenticated and a mentor or service listing is visible. Main flow: learner selects a listing, submits a booking request, the sessions API validates payload and creates a pending session, the mentor receives the pending request, and both users later access room or chat surfaces. Postcondition: a persistent session record exists with a traceable status.", "Normal")],
        "2.3 ER Model": [("p", "UniSkill stores identity, booking, communication, and value exchange in a relational schema. The following ER diagram captures the core entities used by the current system.", "Normal"), ("img", ASSETS / "erd.png", 6.4, "Figure: UniSkill entity relationship diagram"), ("p", "Profiles act as the central entity. Sessions connect mentors and learners, transactions preserve wallet history, reviews capture trust signals, and messages or notifications support coordination around the learning workflow.", "Normal")],
        "2.4 Data flow Diagram": [("p", "The main data-flow perspective for UniSkill is driven by student inputs entering registration, booking, and wallet processes, each of which writes to a controlled data store.", "Normal"), ("img", ASSETS / "dfd.png", 6.4, "Figure: UniSkill data flow diagram"), ("p", "Process specification summary: P1 Registration and Verification validates identity details and writes profile plus vault data; P2 Session Discovery and Booking reads marketplace state and creates or updates sessions; P3 Wallet and Credit Transfer records balance-affecting events and leaves an auditable transaction trail.", "Normal")],
        "2.5 Control Flow Diagram": [("p", "Control flow in UniSkill emphasizes decision points around booking validity, session state transitions, and wallet mutation safety. The control diagram below focuses on the booking-to-settlement path.", "Normal"), ("img", ASSETS / "booking_sequence.png", 6.4, "Figure: Session booking and control progression"), ("p", "Control specifications: booking requests are only accepted from authenticated users; pending records prevent direct transition to completion; and credit-affecting operations must read persistent session and profile state before confirming balance updates.", "Normal")],
        "2.6 State Transition Diagram": [("p", "UniSkill models the session lifecycle explicitly so the interface, notifications, and audit behavior remain consistent across learners and mentors.", "Normal"), ("img", ASSETS / "state_transition.png", 6.4, "Figure: UniSkill session state transition diagram"), ("p", "The important transitions are pending to confirmed, confirmed to in-progress, and in-progress to completed. Cancellation and dispute states branch from mid-lifecycle checkpoints so that the system can preserve recovery and moderation logic.", "Normal")],
        "3.1.1 System Architectural Diagram": [("img", ASSETS / "architecture_overview.png", 6.4, "Figure: UniSkill architectural design")],
        "3.1.2 Description of Architectural Diagram": [("p", "UniSkill is organized as a layered web system. Students interact with the Next.js frontend, which exposes onboarding, profile, explore, wallet, dashboard, and room surfaces. Business actions reach app routes or the Express plus Socket.io backend, where validation, OTP logic, sessions logic, wallet rules, and realtime signaling are applied. Supabase supplies authentication, relational storage, row-level security, storage, and protected session state.", "Normal"), ("p", "This separation keeps public browsing, protected identity, business workflows, and persistence concerns distinct while still supporting a smooth campus-user experience.", "Normal")],
        "3.2.1 Data Dictionary": [("p", "The table below summarizes the most important data structures used by UniSkill.", "Normal"), ("table", ["Data Store / Field Group", "Meaning", "Used By"], [
            ["profiles", "Public user identity, display, wallet snapshot, verification flags", "Signup, profile, dashboard, wallet, reviews"],
            ["user_verification_vault", "Sensitive institutional identity details", "OTP verification and private trust checks"],
            ["sessions", "Booked learning interaction with lifecycle status", "Explore, sessions API, dashboard, room"],
            ["transactions", "Wallet movement and settlement history", "Wallet page, backend validation, audit review"],
            ["service_listings", "Mentor offer with title, category, duration and credit price", "Explore and mentor listing flows"],
            ["messages / notifications", "Coordination and event awareness records", "Room, chat, dashboard and follow-up flows"],
        ])],
        "3.2.1 Normalization": [("p", "The UniSkill schema is normalized to reduce duplication across public identity, private verification, booking, and transaction concerns. Profile data is separated from the verification vault so sensitive identity details do not need to be repeated on public surfaces. Sessions and transactions are stored in dedicated tables rather than embedded inside profiles, which prevents update anomalies and supports auditable historical records.", "Normal"), ("p", "In practical terms the design aligns with first normal form through atomic columns, second normal form through dedicated key-based tables, and third normal form by separating non-key attributes such as reviews, notifications, and messages into their own relations.", "Normal")],
        "3.3.1 Flowchart": [("p", "The following flowchart illustrates a typical mentor onboarding path from account creation to marketplace visibility.", "Normal"), ("img", ASSETS / "mentor_onboarding.png", 5.4, "Figure: Mentor onboarding flowchart")],
        "3.4.1 Screen Shots": [("p", "Screenshots from the running UniSkill frontend are included below to show the current interface treatment for public and onboarding flows.", "Normal"), ("img", SCREENSHOTS / "home.png", 3.3, "Screenshot: UniSkill landing page"), ("img", SCREENSHOTS / "login.png", 3.3, "Screenshot: UniSkill login page"), ("img", SCREENSHOTS / "signup.png", 3.3, "Screenshot: UniSkill signup flow")],
        "Appendix B1: Glossary": [("bullet", [
            "Learner: student seeking help or guidance from a campus peer.",
            "Mentor: student offering a skill, session, or service listing.",
            "Service listing: mentor-authored marketplace offer used for discovery and booking.",
            "Verification vault: private data store for sensitive institutional identity attributes.",
            "Session: structured learning interaction with a persisted lifecycle state.",
            "Credits: internal value units tracked through the wallet and transactions tables.",
        ])],
        "Appendix B2: To Be Determined List": [("bullet", [
            "TBD-1: Final names and roll numbers of the submitting student team.",
            "TBD-2: Final production deployment domain and infrastructure owner.",
            "TBD-3: Whether additional admin moderation screens will be completed in the same academic release.",
        ])],
    }

    for heading, blocks in section_blocks.items():
        start = find_first(doc, heading)
        idx = find_index(doc, heading)
        end = None
        for candidate in doc.paragraphs[idx + 1 :]:
            if is_heading(candidate) or candidate.text.strip().startswith(("2.", "3.", "Appendix")):
                end = candidate
                break
        if end is None:
            # append to end
            end = doc.add_paragraph()
        clear_between(doc, heading, end.text.strip())
        insert_blocks_before(doc, end, blocks)

    remove_instructional_paragraphs(doc)
    trim_after_first_prefix(doc, "TBD-1:")
    trim_body_after_paragraph_prefix(doc, "TBD-1:")

    out = FINAL_DIR / "UNISKILL_SDS_Template_Filled.docx"
    doc.save(out)
    return out


def add_report_chapter(doc: Document, title: str, paragraphs: list[str], images: list[tuple[Path, str]] | None = None, table: tuple[list[str], list[list[str]]] | None = None):
    doc.add_page_break()
    chapter = doc.add_paragraph()
    chapter.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = chapter.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph("")
    if table:
        headers, rows = table
        tbl = doc.add_table(rows=1, cols=len(headers))
        try:
            tbl.style = "Table Grid"
        except KeyError:
            pass
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
        for row in rows:
            cells = tbl.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value
        doc.add_paragraph("")

    for text in paragraphs:
        doc.add_paragraph(text)

    if images:
        for image, caption in images:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image), width=Inches(6.2))
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cp.runs:
                cp.runs[0].italic = True


def build_report() -> Path:
    doc = Document(TEMPLATE_REPORT)
    replace_all(doc, "<TITLE OF THE PROJECT>", PROJECT_LONG_TITLE.upper())
    replace_all(doc, "<Name of the Student(s)>", STUDENT_PLACEHOLDER)
    replace_all(doc, "<University Roll No(s) of the Student(s)>", ROLL_PLACEHOLDER)
    replace_all(doc, "<Name of the Guide(s)>", GUIDE_PLACEHOLDER)
    replace_all(doc, "<MONTH YEAR>", MONTH_YEAR)
    replace_all(doc, "APRIL-2026", MONTH_YEAR)
    replace_all(doc, "<Name of Department>", "Computer Science & Engineering")
    replace_all(doc, "<Name of the Department>", "Computer Science & Engineering")
    replace_all(doc, "<Name, designation and department of the Guide(s)>", "Guide name, designation, and department to be inserted")
    replace_all(doc, "<Name of Internal Guide>", GUIDE_PLACEHOLDER)
    replace_all(doc, "<Name of External Guide (If any)>", "External guide, if applicable")
    replace_all(doc, "Students may write as per their experience.", "This acknowledgement section may be personalized by the submitting team before final submission.")

    # Fill abstract and keywords.
    abstract_paras = [
        "UniSkill is a campus-focused student skill exchange platform designed to help learners discover verified peers, book mentorship or practice sessions, transfer credits through a controlled wallet model, and continue collaboration inside a realtime room experience. The project combines a polished Next.js frontend, Express plus Socket.io service logic, and Supabase-backed storage, authentication, and row-level security controls.",
        "This report presents the motivation, analysis, requirements, architecture, database design, implementation strategy, screenshots, and evaluation perspective for UniSkill. The work positions the platform as both an academic software engineering project and a practical student productivity system that can improve trust, discoverability, and accountability in peer learning.",
    ]
    for idx, text in enumerate(abstract_paras, start=140):
        doc.paragraphs[idx - 1].text = text
    doc.paragraphs[146].text = "Keywords: UniSkill, student skill exchange, peer learning, mentoring marketplace, wallet credits, Supabase, Next.js"

    # Remove template guideline body and replace with project chapters.
    for p in list(doc.paragraphs[211:]):
        delete_paragraph(p)

    syn = extract_sections_by_h1(SOURCE_SYNOPSIS)
    srs = extract_sections_by_h1(SOURCE_SRS)
    uml = extract_sections_by_h1(SOURCE_UML)

    add_report_chapter(
        doc,
        "Chapter-1\nIntroduction",
        compact(syn["1. Introduction"] + syn["3. Problem Definition"][:6]),
        images=[(SCREENSHOTS / "home.png", "Figure 1.1 UniSkill landing experience")],
    )
    add_report_chapter(
        doc,
        "Chapter-2\nLiterature Review",
        compact(syn["2. Literature Review"] + syn["Appendix A: Stakeholder, Market, and Rollout Analysis"][:10]),
    )
    add_report_chapter(
        doc,
        "Chapter-3\nProblem Definition and Objectives",
        compact(syn["3. Problem Definition"] + syn["4. Objectives"] + syn["11. Expected Outcomes and Future Enhancement"]),
        table=(
            ["Concern", "Current informal pattern", "UniSkill response"],
            [
                ["Discovery", "Students rely on referrals or scattered groups", "Searchable mentor and service marketplace"],
                ["Trust", "Little formal proof of identity or quality", "Verification vault, ratings, and session history"],
                ["Accountability", "Payments and commitments are hard to track", "Structured sessions, wallet records, and review workflow"],
            ],
        ),
    )
    add_report_chapter(
        doc,
        "Chapter-4\nSystem Analysis and Requirements",
        compact(srs["1. Introduction"] + srs["2. Overall Description"] + srs["3. External Interface Requirements"][:10]),
        table=(
            ["Layer", "Core Responsibility", "Representative Surfaces"],
            [
                ["User layer", "Student-facing interaction", "Landing page, signup, login, profile, explore"],
                ["Service layer", "Business validation", "OTP routes, sessions API, wallet routes, realtime layer"],
                ["Data layer", "Persistent state and access policy", "Profiles, sessions, transactions, listings, messages"],
            ],
        ),
    )
    add_report_chapter(
        doc,
        "Chapter-5\nDesign and Architecture",
        compact(uml["3. Design Model"][:12] + uml["5. Component and Interface Design"] + uml["8. Design Decisions and Future Evolution"][:8]),
        images=[
            (ASSETS / "architecture_overview.png", "Figure 5.1 Logical architecture of UniSkill"),
            (ASSETS / "component_design.png", "Figure 5.2 Component design view"),
        ],
    )
    add_report_chapter(
        doc,
        "Chapter-6\nDatabase and Security Design",
        compact(srs["6. Data Requirements and Data Dictionary"] + srs["7. Security, Privacy, and Compliance"] + uml["4. Database Design"]),
        images=[
            (ASSETS / "erd.png", "Figure 6.1 Core entity relationship diagram"),
            (ASSETS / "state_transition.png", "Figure 6.2 Session state transition view"),
        ],
        table=(
            ["Data Domain", "Key Fields", "Security Note"],
            [
                ["profiles", "display_name, skills, credits_balance, verification flags", "Public and semi-private user state"],
                ["user_verification_vault", "college_email, phone_number, identity_document_url", "Restricted identity data"],
                ["sessions", "mentor_id, student_id, title, status, room_id", "Lifecycle controlled by authenticated flows"],
                ["transactions", "user_id, session_id, amount, type, status", "Auditable credit movement"],
            ],
        ),
    )
    add_report_chapter(
        doc,
        "Chapter-7\nImplementation and Interface Design",
        compact(syn["8. Module Description"] + syn["7. Methodology and Architecture"] + uml["7. Screen and Interaction Notes"]),
        images=[
            (SCREENSHOTS / "login.png", "Figure 7.1 Login interface"),
            (SCREENSHOTS / "signup.png", "Figure 7.2 Signup interface"),
            (ASSETS / "use_case.png", "Figure 7.3 Use case view"),
        ],
    )
    add_report_chapter(
        doc,
        "Chapter-8\nTesting, Results and Discussion",
        compact(syn["10. Risks, Testing, and Deployment Strategy"] + srs["8. Validation, Acceptance, and Traceability"] + syn["Appendix D: Testing Scenarios and Expected Observations"][:18]),
        images=[
            (ASSETS / "dfd.png", "Figure 8.1 Data flow analysis"),
            (ASSETS / "booking_sequence.png", "Figure 8.2 Booking and settlement sequence"),
        ],
    )
    add_report_chapter(
        doc,
        "Chapter-9\nSummary and Conclusions",
        compact(syn["12. Conclusion and References"] + syn["11. Expected Outcomes and Future Enhancement"][:6] + syn["Appendix C: Demo Flow and Evaluation Readiness"]),
    )

    add_report_chapter(
        doc,
        "Appendix",
        compact(uml["Appendix A: Detailed Entity and Relationship Commentary"][:30] + uml["Appendix B: Use Case Realization Notes"][:25]),
    )
    add_report_chapter(
        doc,
        "Appendix II\nExtended Design Notes",
        compact(uml["Appendix C: Interface-State and Screenflow Commentary"][:70] + uml["Appendix D: Design Alternatives and Trade-off Discussion"][:55]),
        images=[
            (ASSETS / "mentor_onboarding.png", "Figure A2.1 Mentor onboarding flow"),
            (ASSETS / "component_design.png", "Figure A2.2 Component design reference"),
            (ASSETS / "state_transition.png", "Figure A2.3 Session state design reference"),
        ],
    )
    add_report_chapter(
        doc,
        "Appendix III\nRequirement and Recovery Notes",
        compact(srs["Appendix A: Detailed Use Case Specifications"][:40] + srs["Appendix C: Error Handling, Boundary Conditions, and Recovery Requirements"][:35] + srs["Appendix D: Detailed Data Rules and Quality Constraints"][:45]),
        table=(
            ["Scenario", "Failure or boundary condition", "Expected handling"],
            [
                ["Signup verification", "OTP delay or retry threshold reached", "System preserves pending state and exposes safe resend path"],
                ["Booking request", "Invalid slot, balance, or duplicate action", "Request is rejected with audit-safe validation message"],
                ["Wallet settlement", "Repeated completion trigger", "Server prevents duplicate transaction commit"],
                ["Room access", "Unauthorized or stale session state", "Join request is denied until eligibility checks pass"],
            ],
        ),
        images=[
            (ASSETS / "dfd.png", "Figure A3.1 Recovery-sensitive functional flow"),
        ],
    )
    add_report_chapter(
        doc,
        "Appendix IV\nExtended Case and Data Commentary",
        compact(
            uml["Appendix A: Detailed Entity and Relationship Commentary"][30:70]
            + uml["Appendix B: Use Case Realization Notes"][:29]
            + srs["Appendix D: Detailed Data Rules and Quality Constraints"][:40]
        ),
        images=[
            (ASSETS / "erd.png", "Figure A4.1 Extended entity relationship reference"),
            (ASSETS / "booking_sequence.png", "Figure A4.2 Use case realization sequence"),
            (SCREENSHOTS / "home.png", "Figure A4.3 Public platform screenshot"),
        ],
    )
    add_report_chapter(
        doc,
        "Bibliography",
        [
            "1. Next.js Documentation, https://nextjs.org/docs",
            "2. React Documentation, https://react.dev",
            "3. Supabase Documentation, https://supabase.com/docs",
            "4. Express.js Documentation, https://expressjs.com",
            "5. Prisma Documentation, https://docs.prisma.io",
            "6. Socket.IO Documentation, https://socket.io",
            "7. Internal UniSkill references: frontend and backend package manifests, route files, and Supabase migration scripts.",
        ],
    )

    remove_instructional_paragraphs(doc, ["Remove after finalization"])

    out = FINAL_DIR / "UNISKILL_Mini_Project_Report.docx"
    doc.save(out)
    return out


def main() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [
        build_synopsis(),
        build_srs(),
        build_sds(),
        build_report(),
    ]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
