from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def normalize_cover(doc: Document) -> None:
    cover_lines = [
        "UNISKILL",
        "Software Requirement Specification",
        "Systematic and project-specific SRS for the UniSkill student skill exchange platform",
        "",
        "Document Type: Software Requirement Specification",
        "Prepared for final year project submission",
        "Source baseline: UniSkill repository, project documents, and verified product flows",
        "This edition removes unrelated project content and organizes the document for concise academic review",
    ]
    for idx, line in enumerate(cover_lines):
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = line


def add_appendix_expansion(doc: Document) -> None:
    doc.add_page_break()

    heading = doc.add_paragraph("9.3 Pilot Launch and Rollout Notes")
    heading.style = "Heading 2"
    rollout_paragraphs = [
        "UniSkill should be introduced first as a controlled campus pilot for one department or student club so that onboarding quality, wallet behavior, and session coordination can be observed before the platform is opened more broadly.",
        "The first rollout phase should focus on high-signal categories such as programming, design tools, spoken English, aptitude preparation, and resume review because these skills are easy for students to describe, compare, and validate through quick sessions.",
        "A pilot launch must include clear governance rules for verified identity, acceptable service behavior, session cancellation, and review etiquette. These controls reduce the trust deficit that usually weakens informal peer-learning groups.",
        "Awareness should be built through orientation sessions, department groups, and student ambassador campaigns. The goal is not only sign-ups, but also a balanced supply of mentors and learners so the marketplace does not feel empty during its first weeks.",
        "Operational review checkpoints should be scheduled weekly during the pilot. The project team should inspect failed OTP attempts, abandoned bookings, unresolved disputes, room join failures, and wallet anomalies to decide whether the next rollout phase is justified.",
    ]
    for text in rollout_paragraphs:
        doc.add_paragraph(text)

    heading = doc.add_paragraph("9.4 Detailed Module Walkthrough Summary")
    heading.style = "Heading 2"
    module_paragraphs = [
        "Authentication and Verification Module: This module captures the initial student account, checks ownership of the institutional contact method through OTP, and separates public profile data from sensitive verification data stored in the private vault.",
        "Profile and Identity Module: After registration, the student prepares a visible profile with headline, bio, skills, learning goals, pricing signals, and availability indicators so that discovery quality improves across the marketplace.",
        "Marketplace and Discovery Module: The explore surface combines category filters, search, service cards, and mentor summaries so a learner can move quickly from browsing to deciding which peer is most suitable for a session.",
        "Session Lifecycle Module: Booking, approval, scheduling, cancellation, completion, and room access are treated as explicit states. This avoids ambiguity and gives the platform enough structure to support fair review and credit transfer rules.",
        "Wallet and Transaction Module: Credit balance, top-up history, transfer records, and booking-linked deductions must remain consistent across every transaction. The design assumes server-side enforcement for all balance mutations.",
        "Realtime Communication Module: Chat, room signaling, and session-side interaction are not separate add-ons; they are part of the core teaching workflow. This lets UniSkill support the full mentor-learner journey inside one platform.",
        "Governance and Audit Module: Audit events, moderation hooks, verification status, and notification history together provide the accountability layer expected in an academic environment where safety and traceability matter.",
    ]
    for text in module_paragraphs:
        doc.add_paragraph(text)

    heading = doc.add_paragraph("9.5 Success Metrics and Operational Checkpoints")
    heading.style = "Heading 2"
    doc.add_paragraph(
        "The synopsis and design material indicate that UniSkill should be measured as both a learning marketplace and a controlled campus service. The following checkpoints provide a concise way to evaluate whether the first deployment is healthy."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Metric Area", "Primary Measure", "Expected Operational Signal"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    rows = [
        ("Adoption", "Verified student sign-ups and active weekly users", "New users continue beyond signup and create usable profiles"),
        ("Liquidity", "Ratio of published services to learner requests", "Learners can discover multiple relevant mentors per category"),
        ("Trust", "OTP completion rate and dispute frequency", "Most sessions proceed without identity or conduct concerns"),
        ("Execution", "Booking-to-completion ratio and room join success", "Scheduled sessions convert into completed learning events"),
        ("Economy", "Wallet transfer accuracy and refund consistency", "No unexplained balance drift or duplicate transactions"),
        ("Retention", "Repeat sessions and review submission rate", "Students return for additional exchanges and leave feedback"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text

    closing_paragraphs = [
        "If these measures remain healthy during the pilot, the system is ready for broader rollout with stronger moderation tools, richer recommendations, and deeper departmental collaboration.",
        "If the measures show weak marketplace liquidity or repeated session failures, the project should prioritize onboarding quality, mentor seeding, scheduling clarity, and operational support before further expansion.",
    ]
    for text in closing_paragraphs:
        doc.add_paragraph(text)

    doc.add_page_break()
    heading = doc.add_paragraph("9.6 Academic Evaluation Checklist")
    heading.style = "Heading 2"
    checklist_intro = (
        "The SRS, synopsis, and UML-oriented design material all point to the same evaluation logic: UniSkill should be judged not only on interface completeness, but also on whether the platform can support a trustworthy, traceable, and repeatable peer-learning workflow."
    )
    doc.add_paragraph(checklist_intro)

    evaluation_points = [
        "Requirement completeness: evaluators should confirm that signup, verification, profile management, discovery, booking, credits, notifications, and live-room support are all represented clearly in the implemented product.",
        "Consistency across layers: the documented user journey should match the available frontend pages, route handlers, database entities, and policy rules rather than existing only as a conceptual diagram.",
        "Trust and safety posture: the project should demonstrate that sensitive identity data is separated from public data, that wallet mutations are controlled, and that moderation or audit hooks exist for later expansion.",
        "Usability evidence: onboarding steps, mentor discovery, service comparison, and booking confirmation should be understandable to a student user without requiring expert guidance.",
        "Operational readiness: the team should be able to explain deployment assumptions, environment configuration, core dependencies, and the limits of the current pilot release.",
        "Future extensibility: the current version should leave clear space for institutional administration, richer recommendation logic, analytics, and stronger dispute handling without forcing a redesign of the core data model.",
    ]
    for text in evaluation_points:
        doc.add_paragraph(text)

    doc.add_paragraph(
        "This final appendix keeps the first document aligned with typical faculty review expectations while staying tightly focused on UniSkill rather than the unrelated SQL optimizer content that originally occupied the target file."
    )


def clean_document(source: Path, output: Path) -> None:
    doc = Document(source)

    exact_remove = {
        "Priority for this feature group is high because it contributes directly to the core exchange loop of trust, discovery, booking, session execution, or value transfer.",
        "These requirements reflect observed repository behavior and the logical necessities of a campus learning marketplace. Together they define the minimum coherent operating surface for UniSkill.",
        "The implementation evidence in the repository supports these assumptions, but production hardening should make each dependency observable and testable in operations.",
    }

    prefixes = (
        "UniSkill approaches ",
    )

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if text in exact_remove or any(text.startswith(prefix) for prefix in prefixes):
            delete_paragraph(paragraph)

    normalize_cover(doc)
    add_appendix_expansion(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: clean_srs_doc.py <source.docx> <output.docx>")
        return 1

    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    clean_document(source, output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
