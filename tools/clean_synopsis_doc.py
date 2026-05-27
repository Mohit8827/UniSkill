from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def normalize_cover(doc: Document) -> None:
    cover_lines = [
        "UNISKILL",
        "Project Synopsis",
        "A concise academic synopsis for the UniSkill student skill-exchange platform",
        "",
        "Document Type: Synopsis Report",
        "Prepared for final year project submission",
        "Source baseline: UniSkill repository, product flows, and project design documents",
        "This edition removes unrelated project content and presents the UniSkill synopsis in a systematic form",
    ]
    for idx, line in enumerate(cover_lines):
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = line


def add_appendix_expansion(doc: Document) -> None:
    doc.add_page_break()

    heading = doc.add_paragraph("Appendix C: Demo Flow and Evaluation Readiness")
    heading.style = "Heading 1"
    intro = (
        "This appendix summarizes how the UniSkill project can be demonstrated in an academic review setting. The goal is to show that the product is not only conceptually designed, but also organized around a believable end-to-end student workflow."
    )
    doc.add_paragraph(intro)

    subheading = doc.add_paragraph("C.1 Recommended live demonstration flow")
    subheading.style = "Heading 2"
    demo_steps = [
        "Begin with the landing experience and explain the campus-first value proposition of UniSkill: verified students exchanging skills, mentoring, and guided sessions inside one controlled platform.",
        "Show student signup, OTP verification, and profile creation. This step establishes the trust model and explains why identity verification is central to the project.",
        "Move into marketplace discovery where a learner filters profiles, inspects mentor offerings, and selects a relevant service. This proves that the project solves the discovery problem defined earlier in the synopsis.",
        "Demonstrate booking confirmation, wallet interaction, and session lifecycle state changes so evaluators can see how value exchange and accountability are tied together.",
        "Finish with room participation, feedback submission, and notification behavior to show that UniSkill supports the complete mentoring loop rather than stopping at simple profile browsing.",
    ]
    for text in demo_steps:
        doc.add_paragraph(text)

    subheading = doc.add_paragraph("C.2 Evaluation checklist for synopsis presentation")
    subheading.style = "Heading 2"
    doc.add_paragraph(
        "Faculty review of the synopsis usually depends on whether the team can connect the problem statement, objectives, architecture, modules, and expected outcomes into one coherent explanation. The following compact matrix helps structure that discussion."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Evaluation Area", "What To Demonstrate", "Expected Signal"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    rows = [
        ("Problem fit", "Need for verified peer learning and structured campus discovery", "The project solves a real student coordination gap"),
        ("Technical depth", "Frontend, backend, database, and security choices", "The team understands implementation trade-offs"),
        ("Workflow completeness", "Signup through session completion", "UniSkill supports an end-to-end exchange loop"),
        ("Trust model", "OTP, vault design, auditability, and wallet controls", "Safety and accountability are treated seriously"),
        ("Future scope", "Admin governance, analytics, and richer matching", "The project can grow without losing its core architecture"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text

    closing = (
        "Adding this appendix keeps the synopsis aligned with real project presentation needs while staying fully focused on UniSkill instead of the unrelated SQL optimizer content found in the original target file."
    )
    doc.add_paragraph(closing)

    doc.add_page_break()
    subheading = doc.add_paragraph("C.3 Final synopsis takeaway")
    subheading.style = "Heading 2"
    takeaway_paragraphs = [
        "The strongest closing message for the synopsis is that UniSkill is not just a marketplace idea, but a structured campus system that combines discovery, trust, session control, and credit accountability into one coherent workflow.",
        "This framing helps evaluators see the project as both technically grounded and socially relevant: it addresses a real student need while demonstrating modern full-stack engineering, data security awareness, and platform design discipline.",
        "By the end of the presentation, the audience should clearly understand the problem, the proposed solution, the major modules, the implementation approach, and the practical value of deploying UniSkill inside a university environment.",
    ]
    for text in takeaway_paragraphs:
        doc.add_paragraph(text)


def clean_document(source: Path, output: Path) -> None:
    doc = Document(source)

    limited_repeats = {
        "Each of these concerns is directly relevant to UniSkill because the product only succeeds when trust, session completion, and balance integrity work together as one operational chain.": 1,
    }
    repeat_counts: dict[str, int] = {}
    keep_first_uni_paragraph = False

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            keep_first_uni_paragraph = False
            continue

        if text.startswith("UniSkill approaches "):
            if keep_first_uni_paragraph:
                delete_paragraph(paragraph)
                continue
            keep_first_uni_paragraph = True
            continue

        keep_first_uni_paragraph = False

        if text in limited_repeats:
            repeat_counts[text] = repeat_counts.get(text, 0) + 1
            if repeat_counts[text] > limited_repeats[text]:
                delete_paragraph(paragraph)
                continue

        if text.startswith("Figure ") and "High-level logical architecture of UniSkill" in text:
            delete_paragraph(paragraph)

    normalize_cover(doc)
    add_appendix_expansion(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: clean_synopsis_doc.py <source.docx> <output.docx>")
        return 1

    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    clean_document(source, output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
