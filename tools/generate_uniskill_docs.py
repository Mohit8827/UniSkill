from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "generated_docs"
ASSET_DIR = OUT_DIR / "assets"

PRIMARY = RGBColor(31, 64, 104)
ACCENT = RGBColor(221, 110, 70)
SOFT = RGBColor(97, 121, 147)


PROJECT_FACTS = {
    "name": "UniSkill",
    "title": "UniSkill: Verified Campus Skill Exchange Marketplace",
    "subtitle": "Comprehensive project documentation for peer-to-peer academic and extracurricular learning",
    "frontend": "Next.js 16, React 19, Tailwind CSS 4, Framer Motion, Three.js",
    "backend": "Node.js, Express 5, TypeScript, Socket.io",
    "database": "Supabase PostgreSQL with Row Level Security, triggers, and compatibility views",
    "key_modules": [
        "OTP-driven onboarding and university identity verification",
        "Profile, learning goals, and teaching service management",
        "Marketplace discovery for student mentors and campus peers",
        "Session booking with room provisioning and lifecycle tracking",
        "Credit wallet, transaction history, and peer transfers",
        "Messaging, notifications, and WebRTC signaling",
        "Supabase security hardening, audit logging, and verification vault design",
    ],
}


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False, color: RGBColor | None = None, name: str = "Calibri"):
    return {"size": Pt(size), "bold": bold, "color": color, "name": name}


def style_run(run, *, size=11, bold=False, italic=False, color: RGBColor | None = None, name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    if color:
        run.font.color.rgb = color


def set_cell_text(cell, text: str, *, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    style_run(r, size=size, bold=bold)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_doc(doc: Document, header_title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for name, size, color in [
        ("Title", 24, PRIMARY),
        ("Heading 1", 16, PRIMARY),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11, SOFT),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run(header_title)
    style_run(r, size=9, color=SOFT)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("UniSkill Documentation Suite")
    style_run(r, size=9, color=SOFT)
    r.add_break()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def add_cover(doc: Document, title: str, subtitle: str, doc_label: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    r = p.add_run(PROJECT_FACTS["name"].upper())
    style_run(r, size=30, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    style_run(r, size=22, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    style_run(r, size=13, color=SOFT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(24)
    lines = [
        f"Document Type: {doc_label}",
        "Prepared for final year project submission",
        "Source baseline: UniSkill repository, backend routes, frontend flows, and Supabase schema migrations",
        "Generated as a comprehensive reference package with analysis, requirements, design, tables, and diagrams",
    ]
    for idx, line in enumerate(lines):
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx == 0:
            rp.space_before = Pt(18)
        rr = rp.add_run(line)
        style_run(rr, size=12 if idx == 0 else 11, bold=idx == 0, color=PRIMARY if idx == 0 else None)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, first_line_indent=True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.22)
    r = p.add_run(text)
    style_run(r, size=11)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r, size=10)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r, size=10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, head in enumerate(headers):
        set_cell_text(hdr[idx], head, bold=True, size=10)
        shade_cell(hdr[idx], "DDEBF7")
        if widths:
            hdr[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def generate_paragraph(topic: str, angle: str, detail: str, lens: str) -> str:
    return (
        f"UniSkill approaches {topic} through a campus-first design lens in which {angle}. "
        f"The current repository shows that the platform blends {detail}, allowing the system to support both immediate usability and disciplined back-end control. "
        f"From a project documentation perspective, this matters because {lens}, which makes the design traceable from user need through schema, API behavior, and interface flow."
    )


def build_filler_block(doc: Document, topic: str, variants: list[tuple[str, str, str]]) -> None:
    for angle, detail, lens in variants:
        add_para(doc, generate_paragraph(topic, angle, detail, lens))


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    style_run(r, size=10, italic=True, color=SOFT)


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, fill, outline, font_obj):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    w = x2 - x1
    h = y2 - y1
    lines = text.split("\n")
    line_h = 22
    total = len(lines) * line_h
    y = y1 + (h - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font_obj)
        lw = bbox[2] - bbox[0]
        draw.text((x1 + (w - lw) / 2, y), line, fill=(28, 39, 53), font=font_obj)
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, start, end, fill=(80, 92, 108), width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    left = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=fill)


def diagram_canvas(name: str, size=(1600, 900)) -> tuple[Image.Image, ImageDraw.ImageDraw, Path]:
    ensure_dir(ASSET_DIR)
    path = ASSET_DIR / f"{name}.png"
    img = Image.new("RGB", size, (248, 250, 252))
    draw = ImageDraw.Draw(img)
    return img, draw, path


def get_font(size: int):
    try:
        return ImageFont.truetype("arial.ttf", size)
    except Exception:
        return ImageFont.load_default()


def create_architecture_diagram() -> Path:
    img, draw, path = diagram_canvas("architecture_overview")
    title_font = get_font(30)
    box_font = get_font(24)
    draw.text((60, 40), "UniSkill Logical Architecture", fill=(31, 64, 104), font=title_font)
    draw_box(draw, (80, 150, 430, 330), "Students\nLearners / Mentors / Admin", (255, 244, 238), (221, 110, 70), box_font)
    draw_box(draw, (520, 120, 910, 360), "Next.js Frontend\nAuthContext\nMarketplace\nProfile\nWallet\nRoom UI", (230, 239, 250), (31, 64, 104), box_font)
    draw_box(draw, (1010, 120, 1490, 360), "Express + Socket.io Backend\nOTP Routes\nAuth Routes\nWallet Routes\nSkills Match API\nRealtime Signaling", (233, 245, 239), (85, 127, 110), box_font)
    draw_box(draw, (520, 500, 910, 760), "Supabase Platform\nAuth\nRLS Policies\nPostgreSQL Tables\nStorage\nSSR Session Cookies", (244, 241, 255), (115, 97, 163), box_font)
    draw_box(draw, (1010, 500, 1490, 760), "Core Data Domains\nProfiles\nSessions\nTransactions\nReviews\nMessages\nNotifications\nService Listings", (255, 247, 230), (168, 125, 54), box_font)
    arrow(draw, (430, 240), (520, 240))
    arrow(draw, (910, 240), (1010, 240))
    arrow(draw, (715, 360), (715, 500))
    arrow(draw, (1250, 360), (1250, 500))
    arrow(draw, (910, 620), (1010, 620))
    img.save(path)
    return path


def create_use_case_diagram() -> Path:
    img, draw, path = diagram_canvas("use_case")
    title_font = get_font(30)
    box_font = get_font(22)
    draw.text((60, 35), "UniSkill Use Case View", fill=(31, 64, 104), font=title_font)
    draw.ellipse((70, 220, 190, 420), outline=(221, 110, 70), width=4, fill=(255, 245, 240))
    draw.ellipse((70, 500, 190, 700), outline=(31, 64, 104), width=4, fill=(235, 242, 251))
    draw.ellipse((1380, 360, 1510, 560), outline=(85, 127, 110), width=4, fill=(234, 245, 239))
    draw.text((103, 305), "Learner", fill=(28, 39, 53), font=box_font)
    draw.text((108, 585), "Mentor", fill=(28, 39, 53), font=box_font)
    draw.text((1409, 445), "Admin", fill=(28, 39, 53), font=box_font)
    cases = [
        (500, 140, "Register & Verify"),
        (500, 280, "Manage Profile"),
        (500, 420, "Browse Marketplace"),
        (500, 560, "Book Session"),
        (900, 210, "Create Service Listing"),
        (900, 350, "Chat / Join Room"),
        (900, 490, "Transfer Credits"),
        (900, 630, "Rate / Review"),
        (1200, 280, "Moderate Trust Signals"),
        (1200, 520, "Audit Disputes"),
    ]
    centers = []
    for x, y, label in cases:
        draw.ellipse((x, y, x + 240, y + 90), outline=(97, 121, 147), width=3, fill=(255, 255, 255))
        bbox = draw.textbbox((0, 0), label, font=box_font)
        draw.text((x + 120 - (bbox[2] - bbox[0]) / 2, y + 28), label, fill=(28, 39, 53), font=box_font)
        centers.append((x + 120, y + 45))
    for target in centers[:4]:
        arrow(draw, (190, 320), target)
    for target in centers[1:8]:
        arrow(draw, (190, 600), target)
    for target in centers[8:]:
        arrow(draw, (1380, 460), target)
    img.save(path)
    return path


def create_sequence_diagram() -> Path:
    img, draw, path = diagram_canvas("booking_sequence", size=(1600, 1100))
    title_font = get_font(30)
    font_obj = get_font(20)
    draw.text((60, 35), "Session Booking and Credit Reservation Sequence", fill=(31, 64, 104), font=title_font)
    actors = ["Learner", "Frontend", "Sessions API", "Supabase", "Mentor"]
    xs = [140, 430, 740, 1060, 1360]
    for x, actor in zip(xs, actors):
        draw_box(draw, (x - 80, 100, x + 80, 170), actor, (255, 255, 255), (97, 121, 147), font_obj)
        draw.line((x, 170, x, 1000), fill=(180, 188, 197), width=2)
    steps = [
        (220, 140, 430, "Open mentor card and submit booking"),
        (300, 430, 740, "POST /api/sessions"),
        (380, 740, 1060, "Insert pending session + room_id"),
        (460, 1060, 740, "Session row returned"),
        (540, 740, 1060, "Read profile / wallet state"),
        (620, 740, 1360, "Trigger mentor notification"),
        (700, 740, 430, "Return success payload"),
        (780, 430, 140, "Render pending confirmation"),
        (860, 1360, 430, "Mentor later joins chat / room"),
    ]
    for y, x1, x2, text in steps:
        arrow(draw, (x1, y), (x2, y))
        draw.text((min(x1, x2) + 20, y - 30), text, fill=(28, 39, 53), font=font_obj)
    img.save(path)
    return path


def create_erd_diagram() -> Path:
    img, draw, path = diagram_canvas("erd", size=(1700, 1200))
    title_font = get_font(30)
    box_font = get_font(20)
    draw.text((60, 35), "Entity Relationship Diagram", fill=(31, 64, 104), font=title_font)
    entities = {
        "profiles": ((90, 140, 430, 420), ["id (PK)", "name", "email", "credits", "is_mentor", "rating"]),
        "sessions": ((520, 120, 900, 430), ["id (PK)", "mentor_id (FK)", "student_id (FK)", "title", "status", "room_id"]),
        "transactions": ((1010, 140, 1390, 430), ["id (PK)", "user_id (FK)", "session_id (FK)", "amount", "type", "status"]),
        "reviews": ((90, 560, 430, 860), ["id (PK)", "session_id (FK)", "reviewer_id (FK)", "reviewee_id (FK)", "rating", "comment"]),
        "messages": ((520, 560, 900, 860), ["id (PK)", "sender_id (FK)", "receiver_id (FK)", "content", "is_read", "created_at"]),
        "service_listings": ((1010, 560, 1470, 900), ["id (PK)", "mentor_id (FK)", "title", "category", "duration_minutes", "price_credits"]),
    }
    for name, (box, fields) in entities.items():
        draw.rounded_rectangle(box, radius=20, fill=(255, 255, 255), outline=(97, 121, 147), width=3)
        x1, y1, x2, y2 = box
        draw.rectangle((x1, y1, x2, y1 + 46), fill=(221, 235, 247), outline=(97, 121, 147))
        draw.text((x1 + 12, y1 + 10), name, fill=(31, 64, 104), font=box_font)
        y = y1 + 62
        for field in fields:
            draw.text((x1 + 12, y), field, fill=(28, 39, 53), font=box_font)
            y += 34
    for pair in [((430, 250), (520, 250)), ((900, 250), (1010, 250)), ((260, 420), (260, 560)), ((710, 430), (710, 560)), ((1110, 430), (1180, 560)), ((430, 250), (90, 700)), ((430, 250), (520, 700))]:
        arrow(draw, pair[0], pair[1])
    img.save(path)
    return path


def create_dfd_diagram() -> Path:
    img, draw, path = diagram_canvas("dfd")
    title_font = get_font(30)
    box_font = get_font(22)
    draw.text((60, 35), "Data Flow Diagram", fill=(31, 64, 104), font=title_font)
    draw_box(draw, (90, 280, 320, 470), "Student User", (255, 245, 240), (221, 110, 70), box_font)
    draw_box(draw, (460, 150, 810, 340), "P1 Registration & Verification", (230, 239, 250), (31, 64, 104), box_font)
    draw_box(draw, (460, 390, 810, 580), "P2 Session Discovery & Booking", (230, 239, 250), (31, 64, 104), box_font)
    draw_box(draw, (460, 630, 810, 820), "P3 Wallet & Credit Transfer", (230, 239, 250), (31, 64, 104), box_font)
    draw_box(draw, (980, 180, 1420, 340), "D1 Profiles + Vault", (255, 255, 255), (97, 121, 147), box_font)
    draw_box(draw, (980, 430, 1420, 590), "D2 Sessions + Listings", (255, 255, 255), (97, 121, 147), box_font)
    draw_box(draw, (980, 680, 1420, 840), "D3 Transactions + Audit Log", (255, 255, 255), (97, 121, 147), box_font)
    for a, b in [
        ((320, 350), (460, 245)),
        ((320, 375), (460, 485)),
        ((320, 395), (460, 725)),
        ((810, 245), (980, 245)),
        ((810, 485), (980, 505)),
        ((810, 725), (980, 760)),
    ]:
        arrow(draw, a, b)
    img.save(path)
    return path


def create_state_diagram() -> Path:
    img, draw, path = diagram_canvas("state_transition")
    title_font = get_font(30)
    box_font = get_font(22)
    draw.text((60, 35), "Session State Transition Diagram", fill=(31, 64, 104), font=title_font)
    states = [
        (100, 370, "Requested"),
        (360, 370, "Pending"),
        (620, 370, "Confirmed"),
        (880, 370, "In Progress"),
        (1140, 370, "Completed"),
        (620, 650, "Cancelled"),
        (880, 650, "Disputed"),
    ]
    for x, y, label in states:
        draw.ellipse((x, y, x + 180, y + 110), fill=(255, 255, 255), outline=(97, 121, 147), width=3)
        bbox = draw.textbbox((0, 0), label, font=box_font)
        draw.text((x + 90 - (bbox[2] - bbox[0]) / 2, y + 40), label, fill=(28, 39, 53), font=box_font)
    arrows = [
        ((280, 425), (360, 425)),
        ((540, 425), (620, 425)),
        ((800, 425), (880, 425)),
        ((1060, 425), (1140, 425)),
        ((710, 480), (710, 650)),
        ((970, 480), (970, 650)),
    ]
    for a, b in arrows:
        arrow(draw, a, b)
    img.save(path)
    return path


def create_component_diagram() -> Path:
    img, draw, path = diagram_canvas("component_design")
    title_font = get_font(30)
    box_font = get_font(22)
    draw.text((60, 35), "Component Design View", fill=(31, 64, 104), font=title_font)
    comps = [
        ((90, 140, 420, 300), "Landing and Authentication UI"),
        ((90, 360, 420, 520), "Profile and Service Listing UI"),
        ((90, 580, 420, 740), "Explore, Wallet, Dashboard, Room UI"),
        ((540, 140, 960, 300), "SSR Routes and App API Endpoints"),
        ((540, 360, 960, 520), "Express Business Services"),
        ((540, 580, 960, 740), "Socket.io Realtime Layer"),
        ((1080, 140, 1520, 300), "Supabase Auth and Cookie Session"),
        ((1080, 360, 1520, 520), "Relational Storage and Policies"),
        ((1080, 580, 1520, 740), "Triggers, Audit, Notifications"),
    ]
    for box, text in comps:
        draw_box(draw, box, text, (255, 255, 255), (97, 121, 147), box_font)
    for a, b in [
        ((420, 220), (540, 220)),
        ((420, 440), (540, 440)),
        ((420, 660), (540, 660)),
        ((960, 220), (1080, 220)),
        ((960, 440), (1080, 440)),
        ((960, 660), (1080, 660)),
    ]:
        arrow(draw, a, b)
    img.save(path)
    return path


def create_activity_diagram() -> Path:
    img, draw, path = diagram_canvas("mentor_onboarding")
    title_font = get_font(30)
    box_font = get_font(22)
    draw.text((60, 35), "Mentor Onboarding Activity Diagram", fill=(31, 64, 104), font=title_font)
    steps = [
        ((650, 80, 920, 160), "Start"),
        ((650, 200, 920, 290), "Create account"),
        ((650, 330, 920, 420), "Verify OTP and identity"),
        ((650, 460, 920, 550), "Add teachable skills"),
        ((650, 590, 920, 680), "Publish service listing"),
        ((650, 720, 920, 810), "Visible in marketplace"),
    ]
    for box, label in steps:
        draw_box(draw, box, label, (255, 255, 255), (97, 121, 147), box_font)
    for idx in range(len(steps) - 1):
        x = 785
        arrow(draw, (x, steps[idx][0][3]), (x, steps[idx + 1][0][1]))
    img.save(path)
    return path


def make_diagrams() -> dict[str, Path]:
    return {
        "architecture": create_architecture_diagram(),
        "use_case": create_use_case_diagram(),
        "sequence": create_sequence_diagram(),
        "erd": create_erd_diagram(),
        "dfd": create_dfd_diagram(),
        "state": create_state_diagram(),
        "component": create_component_diagram(),
        "activity": create_activity_diagram(),
    }


def add_manual_toc(doc: Document, items: list[str], heading: str = "Table of Contents") -> None:
    add_heading(doc, heading, 1)
    for item in items:
        p = doc.add_paragraph()
        r = p.add_run(item)
        style_run(r, size=10)
    doc.add_page_break()


def section_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_detail_appendix(doc: Document, heading: str, topics: list[str], *, base_number: str) -> None:
    add_heading(doc, heading, 1)
    for idx, topic in enumerate(topics, start=1):
        add_heading(doc, f"{base_number}.{idx} {topic}", 2)
        build_filler_block(doc, topic, [
            ("the project must translate abstract product intent into operational behavior", "student-facing interactions, server validation, and storage outcomes", "design review can evaluate whether the platform is internally coherent"),
            ("UniSkill behaves as a living service ecosystem rather than a one-time form submission", "identity, discovery, booking, communication, and credit loops", "documentation needs enough depth to describe sustained use"),
            ("each topic affects both usability and platform trust", "visible interface cues and invisible policy or schema enforcement", "engineering and product decisions remain traceable"),
            ("institutional adoption would depend on consistent treatment of this topic", "roles, policies, and structured records", "the report should show implementation maturity beyond a prototype demo"),
        ])
        add_bullets(doc, [
            f"Design note for {topic.lower()}: the feature must remain understandable for a first-time student user.",
            f"Technical note for {topic.lower()}: API, schema, and UI state should remain synchronized.",
            f"Quality note for {topic.lower()}: failures should preserve traceability and user confidence.",
        ])


def add_use_case_spec(doc: Document, code: str, name: str, actor: str, preconditions: list[str], flow: list[str], outcomes: list[str]) -> None:
    add_heading(doc, f"{code} {name}", 2)
    add_table(
        doc,
        ["Field", "Specification"],
        [
            ["Primary actor", actor],
            ["Goal", name],
            ["Preconditions", "; ".join(preconditions)],
            ["Postconditions", "; ".join(outcomes)],
        ],
        widths=[1.5, 4.9],
    )
    add_heading(doc, f"{code}.1 Main Flow", 3)
    add_numbered(doc, flow)
    add_heading(doc, f"{code}.2 Notes", 3)
    build_filler_block(doc, name, [
        ("the scenario must be intelligible from both product and engineering perspectives", "user actions, route execution, and resulting state changes", "use-case specification becomes a bridge between requirements and implementation"),
        ("UniSkill often couples trust decisions with convenience expectations", "verification feedback, booking state, and balance rules", "alternate flows deserve explicit attention"),
    ])


def build_synopsis(diagrams: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc, "UniSkill Project Synopsis")
    add_cover(
        doc,
        "Project Synopsis",
        "A detailed academic synopsis for the UniSkill student skill-exchange platform",
        "Synopsis Report",
    )
    toc = [
        "1. Introduction",
        "2. Literature Review",
        "3. Problem Definition",
        "4. Objectives",
        "5. Scope and Feasibility",
        "6. Existing and Proposed System",
        "7. Methodology and Architecture",
        "8. Module Description",
        "9. Requirements and Tools",
        "10. Risks, Testing, and Deployment Strategy",
        "11. Expected Outcomes and Future Enhancement",
        "12. Conclusion and References",
    ]
    add_manual_toc(doc, toc)

    add_heading(doc, "1. Introduction", 1)
    intro_variants = [
        ("students can be both mentors and learners inside the same platform", "responsive pages, Supabase-backed sessions, and credit-oriented service flows", "the system can be positioned as both a social learning product and an engineering artifact"),
        ("the marketplace is restricted to trusted campus users rather than anonymous internet traffic", "verification steps, profile enrichment, and server-mediated booking", "the trust model becomes a first-class design concern instead of an afterthought"),
        ("its value grows from repeated peer exchange rather than one-off tutoring sales", "wallet tracking, session history, and reviewable mentor offerings", "long-term engagement metrics are as important as simple sign-up counts"),
        ("frontend polish is paired with a fairly rigorous storage and policy layer", "Next.js interaction design together with RLS, audit logging, and OTP records", "the report can explain both usability and governance in a balanced way"),
    ]
    build_filler_block(doc, "student skill exchange", intro_variants)
    add_bullets(doc, [f"Core module: {item}" for item in PROJECT_FACTS["key_modules"]])
    add_para(doc, "The educational context for UniSkill is shaped by uneven access to tutoring, fragmented peer networks, and the absence of an institutional mechanism that lets students transparently monetize or exchange expertise without leaving the campus context.")
    add_figure(doc, diagrams["architecture"], "Figure 1.1: High-level logical architecture of UniSkill", width=6.4)

    add_heading(doc, "1.1 Overview", 2)
    for paragraph in [
        "UniSkill is conceived as a hybrid learning marketplace in which a student can discover another student by skill, reputation, and listed services, initiate a session request, and later move into a live room supported by signaling, chat, and structured status updates.",
        "The repository reveals a multi-surface product rather than a static portal. The landing area introduces the product, the dashboard summarizes activity, the explore area surfaces mentors and service listings, the profile area controls skills and offerings, and route handlers persist the life cycle through Supabase-backed APIs.",
        "This synopsis treats the project as both a deployable web application and an academic systems project, so the discussion covers user value, data integrity, security rules, implementation trade-offs, and measurable project outcomes.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "1.2 Problem Statement", 2)
    add_numbered(doc, [
        "Students usually find peer help through messaging groups, personal referrals, or chance interaction, which makes discovery unreliable and heavily biased toward existing social circles.",
        "Learners often cannot compare teaching quality, availability, session structure, or expected cost before they commit to a session, so trust and decision quality remain low.",
        "Students who can teach valuable skills have no structured campus-native system for building credibility, listing offers, and receiving value in a reusable campus economy.",
        "Universities increasingly encourage collaborative learning, yet operational systems for verified peer exchange, credit flow, and accountable mentorship are usually absent.",
    ])

    add_heading(doc, "2. Literature Review", 1)
    literature_topics = [
        "peer learning ecosystems",
        "two-sided marketplaces",
        "digital identity verification in education",
        "virtual credit systems",
        "real-time collaboration tools",
        "row-level security in SaaS products",
    ]
    for idx, topic in enumerate(literature_topics, start=1):
        add_heading(doc, f"2.{idx} Review Area: {topic.title()}", 2)
        build_filler_block(doc, topic, [
            ("successful platforms reduce friction between intent and action", "search, filtering, visible trust cues, and transactional closure", "user adoption depends on the entire workflow rather than on isolated features"),
            ("a strong data model matters as much as interface convenience", "profiles, offerings, session records, and transaction logs", "analysis can tie literature concepts to concrete tables and routes in UniSkill"),
            ("platform trust grows when verification, transparency, and recovery paths are explicit", "OTP processes, audit-friendly schema, and role-sensitive policies", "the academic justification for the proposed system becomes more defensible"),
        ])
    add_heading(doc, "2.7 Gap Analysis", 2)
    add_para(doc, "Existing tutoring and freelance systems are either too generic, too commercially oriented, or too weak on institutional trust. UniSkill fills the gap by combining campus identity, peer-to-peer service exchange, educational intent, and a reusable credits model inside one coherent architecture.")

    add_heading(doc, "3. Problem Definition", 1)
    for heading, topic in [
        ("3.1 Discovery Friction", "mentor discovery"),
        ("3.2 Trust Deficit", "identity and reliability"),
        ("3.3 Informal Payments and Accountability", "wallet control"),
        ("3.4 Fragmented Communication", "session coordination"),
    ]:
        add_heading(doc, heading, 2)
        build_filler_block(doc, topic, [
            ("current student workflows rely on manual coordination and uncertain follow-through", "search pages, profile metadata, and room identifiers", "the need for structured orchestration becomes obvious"),
            ("the absence of system records weakens both confidence and dispute resolution", "transaction logs, review records, and audit artifacts", "the business rationale for persistent data becomes academically measurable"),
        ])

    add_heading(doc, "4. Objectives", 1)
    objectives = [
        "Build a verified skill exchange platform tailored to university students and lightweight enough for repeated daily use.",
        "Design a modular architecture in which frontend, backend, realtime transport, and Supabase persistence can evolve independently.",
        "Establish a trust framework through OTP verification, private verification vault storage, and role-based data access rules.",
        "Create a sustainable credit economy that rewards teaching, tracks transfers, and preserves a clear financial-style audit trail.",
        "Support mentorship discovery and booking through profile signals, service listings, session creation, and real-time follow-up.",
        "Provide an extensible base for future admin analytics, dispute management, automated matchmaking, and institutional rollout.",
    ]
    add_numbered(doc, objectives)

    add_heading(doc, "5. Scope and Feasibility", 1)
    for sub, items in [
        ("5.1 Scope of Current Implementation", [
            "Account creation, login, OTP dispatch, and verification status tracking",
            "Profile compatibility layer for both legacy and reconciled Supabase schema fields",
            "Marketplace exploration for mentors and service listings",
            "Session record creation and room identifier generation",
            "Wallet viewing, credit purchase simulation, and credit transfer routes",
            "Realtime signaling and chat events through Socket.io rooms",
        ]),
        ("5.2 Operational Feasibility", [
            "Students already use browsers and mobile-responsive portals as their default campus interaction surface.",
            "The system stack runs on mainstream developer tools and does not require unusual local hardware.",
            "Supabase reduces infrastructure friction by combining auth, database, storage, and policy controls.",
        ]),
        ("5.3 Technical Feasibility", [
            "The chosen stack is mature, well-supported, and already represented in the repository.",
            "Data-intensive requirements are moderate, making PostgreSQL and RLS an appropriate fit.",
            "The architecture can be deployed incrementally, allowing staged validation of each module.",
        ]),
        ("5.4 Economic Feasibility", [
            "The platform can begin on low-cost managed services while serving a limited campus cohort.",
            "Virtual credits minimize payment gateway dependence during early validation phases.",
            "Reuse of open-source frameworks reduces capital expenditure for prototyping and iteration.",
        ]),
    ]:
        add_heading(doc, sub, 2)
        add_bullets(doc, items)
        add_para(doc, f"{sub.split()[0]} is strengthened by the fact that UniSkill already demonstrates a working integration pattern across interface, API endpoints, and schema migration artifacts.")

    add_heading(doc, "6. Existing and Proposed System", 1)
    add_heading(doc, "6.1 Existing System", 2)
    add_table(
        doc,
        ["Aspect", "Existing informal practice", "Limitation"],
        [
            ["Discovery", "WhatsApp groups, peers, referrals", "No structured search or rating context"],
            ["Trust", "Personal belief", "No auditable verification or persistent service history"],
            ["Payment", "Cash or favors", "Awkward, inconsistent, and hard to reconcile"],
            ["Scheduling", "Manual chat threads", "High coordination cost and missed follow-up"],
            ["Quality", "Word of mouth", "No unified review or session record"],
        ],
        widths=[1.3, 2.4, 2.6],
    )
    add_heading(doc, "6.2 Proposed System", 2)
    build_filler_block(doc, "the proposed UniSkill system", [
        ("each user journey is stored as a combination of profile, service, session, and transaction state", "server route orchestration and Supabase persistence", "functional scope can be traced to measurable entities"),
        ("verification and privacy are separated instead of mixed into one public profile object", "the verification vault, public profile fields, and SSR auth client", "sensitive data handling becomes explainable and reviewable"),
        ("the system intentionally supports both a modern learner-facing experience and back-office governance potential", "frontend dashboards, audit logging, and role-aware policies", "the project scales from prototype to institutional pilot more naturally"),
    ])

    add_heading(doc, "7. Methodology and Architecture", 1)
    add_heading(doc, "7.1 Development Methodology", 2)
    for text in [
        "An incremental and feedback-driven methodology suits UniSkill because user experience, trust flow, and schema evolution influence one another. The repository shows iterative refinement, including schema reconciliation migrations, compatibility handling inside the frontend auth context, and modular route definition in the backend.",
        "From a project management standpoint, the work can be framed as an agile campus product: onboarding and profile flows validate trust assumptions first, marketplace and service listing flows validate discoverability, and session plus wallet features validate the exchange model.",
    ]:
        add_para(doc, text)
    add_heading(doc, "7.2 System Architecture Overview", 2)
    add_figure(doc, diagrams["component"], "Figure 7.2: Component design spanning UI surfaces, route handlers, and persistence services", width=6.5)
    add_para(doc, "The architecture is layered but not rigidly monolithic. Client pages in the Next.js app own user interaction and SSR-aware session handling. Route handlers provide server-side mediation for profiles, sessions, and service listings. Express routes handle OTP, wallet behavior, auth-specific logic, and realtime event transport. Supabase anchors identity, relational storage, and security policy enforcement.")
    add_heading(doc, "7.3 Database and Security Strategy", 2)
    build_filler_block(doc, "database security", [
        ("public profile discoverability coexists with private verification storage", "profiles, user_verification_vault, and RLS policies", "privacy boundaries can be clearly justified"),
        ("auditability is embedded in the schema rather than postponed to administration tooling", "user_audit_log and trigger-based tracking", "institutional trust discussions remain evidence-based"),
        ("compatibility with legacy fields reduces migration shock during product evolution", "credits versus credits_balance and name versus display_name handling", "the project demonstrates realistic engineering trade-offs"),
    ])

    add_heading(doc, "8. Module Description", 1)
    modules = [
        ("8.1 Authentication and Verification Module", "registration, login, OTP sending, OTP verification, verification status checks, and secure session persistence"),
        ("8.2 Profile and Identity Module", "student-facing profile editing, compatibility mapping, public trust indicators, and teaching or learning metadata"),
        ("8.3 Marketplace and Discovery Module", "mentor exploration, filters, skill matching, service listing retrieval, and campus-oriented categorization"),
        ("8.4 Session Lifecycle Module", "booking initiation, room creation, pending and confirmed states, realtime handoff, and session completion intent"),
        ("8.5 Wallet and Transaction Module", "wallet snapshots, transaction history, credit top-up simulation, and peer transfer logic"),
        ("8.6 Realtime Communication Module", "Socket.io room join events, signal relay, chat message propagation, and call termination notifications"),
        ("8.7 Admin and Governance Extension Module", "verification review, dispute handling, analytics, and trust moderation envisioned over the current base"),
    ]
    for heading, desc in modules:
        add_heading(doc, heading, 2)
        add_para(doc, f"This module covers {desc}. It is central to the overall value proposition because it turns the abstract idea of student skill exchange into a repeatable operational flow.")
        build_filler_block(doc, heading.split(":")[0], [
            ("the module must balance convenience with validation", desc, "implementation quality depends on preserving both speed and control"),
            ("the repository already provides concrete hooks for this behavior", "frontend page flows, route handlers, and persistent tables", "documentation can map design claims to source artifacts"),
            ("future iterations can extend the same module without rewriting the entire platform", "clear ownership boundaries and modular storage structures", "the system remains maintainable as requirements grow"),
        ])

    add_heading(doc, "9. Requirements and Tools", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Usage in UniSkill"],
        [
            ["Frontend", PROJECT_FACTS["frontend"], "Interactive UI, dashboards, profile forms, visual polish, browser rendering"],
            ["Backend", PROJECT_FACTS["backend"], "RESTful business routes, OTP handling, wallet logic, realtime signaling"],
            ["Database", PROJECT_FACTS["database"], "Relational persistence, auth, policies, triggers, and storage"],
            ["ORM / Data Access", "Prisma and Supabase clients", "Typed schema reference plus query and auth operations"],
            ["Utilities", "Nodemailer, SSR cookie helpers, Lucide, motion libraries", "Emails, sessions, icons, and interaction quality"],
        ],
        widths=[1.1, 2.7, 2.5],
    )
    add_para(doc, "Hardware requirements remain modest for development, while production planning should emphasize network reliability, managed database throughput, secure secret storage, and observability around authentication and booking flows.")

    add_heading(doc, "10. Risks, Testing, and Deployment Strategy", 1)
    for sub in [
        ("10.1 Key Risks", [
            "Schema drift between frontend expectations and live database structure",
            "Weak verification if OTP and vault flows are not tightly validated",
            "Credit inconsistencies if transfer operations are not wrapped in stronger transaction discipline",
            "Realtime instability under poor networks or abrupt room exits",
        ]),
        ("10.2 Testing Strategy", [
            "Unit testing for validation helpers and route-level business checks",
            "Integration tests for auth, session creation, and wallet transfer paths",
            "UI scenario tests for signup, profile updates, mentor discovery, and room entry",
            "Security verification of RLS behavior and sensitive-field exposure",
        ]),
        ("10.3 Deployment Strategy", [
            "Deploy frontend and backend separately but with aligned environment variables and origin rules",
            "Use Supabase migrations as the canonical database change log",
            "Monitor auth success, OTP delivery, booking completion, and transfer failure rates from the first pilot cohort",
        ]),
    ]:
        add_heading(doc, sub[0], 2)
        add_bullets(doc, sub[1])
        add_para(doc, "Each of these concerns is directly relevant to UniSkill because the product only succeeds when trust, session completion, and balance integrity work together as one operational chain.")

    add_heading(doc, "11. Expected Outcomes and Future Enhancement", 1)
    build_filler_block(doc, "expected outcomes", [
        ("students gain a clearer path to either seek help or monetize expertise", "discoverable profiles, session records, and credit-based exchange", "impact can be assessed through retention and session activity"),
        ("campus communities gain a reusable marketplace rather than isolated tutoring threads", "structured data, persistent ratings, and repeatable booking flows", "the project supports both educational and entrepreneurial narratives"),
        ("future enhancements can deepen rather than replace the current architecture", "recommendation engines, admin panels, analytics, and dispute resolution tooling", "the system remains a strong base for ongoing research and product work"),
    ])
    add_heading(doc, "11.1 Future Scope", 2)
    add_numbered(doc, [
        "AI-assisted mentor recommendation based on learning goals and prior outcomes.",
        "Institution admin console for verification review, trust intervention, and usage analytics.",
        "Calendar integration, reminders, and richer availability scheduling.",
        "Session evidence capture, refunds, and dispute escalation workflows.",
        "Cross-campus federation with domain-specific verification controls.",
    ])

    add_heading(doc, "12. Conclusion and References", 1)
    add_para(doc, "UniSkill stands out as a practical and academically meaningful software project because it combines modern full-stack engineering with a well-defined social need inside the university ecosystem. Its architecture is credible, its problem framing is concrete, and its implementation already shows the movement from concept to deployable product.")
    add_heading(doc, "12.1 References", 2)
    add_bullets(doc, [
        "UniSkill frontend and backend repository modules.",
        "Supabase migration scripts defining profile, session, transaction, and verification structures.",
        "Next.js, React, Express, Prisma, and Socket.io official documentation for implementation alignment.",
    ])

    add_detail_appendix(
        doc,
        "Appendix A: Stakeholder, Market, and Rollout Analysis",
        [
            "Primary stakeholder expectations",
            "Learner success metrics",
            "Mentor motivation structure",
            "Campus trust and institutional concerns",
            "Adoption barriers in the early pilot phase",
            "Marketplace liquidity strategy",
            "Credit-economy sustainability assumptions",
            "Support and moderation considerations",
            "Communication and onboarding campaign ideas",
            "Rollout sequencing for campus launch",
        ],
        base_number="A",
    )
    add_detail_appendix(
        doc,
        "Appendix B: Detailed Module Walkthroughs and Scenario Narratives",
        [
            "Signup and identity capture scenario",
            "OTP verification scenario",
            "Profile completion and mentor activation",
            "Teaching-service publication lifecycle",
            "Mentor discovery and comparison",
            "Booking confirmation and room preparation",
            "Wallet top-up and transfer interpretation",
            "Realtime room participation flow",
            "Review and feedback closure",
            "Maintenance and future admin intervention points",
        ],
        base_number="B",
    )
    add_heading(doc, "Appendix C: Implementation Timeline and Work Breakdown", 1)
    add_table(
        doc,
        ["Phase", "Focus", "Representative outcomes"],
        [
            ["Phase 1", "Foundation", "Project setup, auth integration, schema baseline, landing experience"],
            ["Phase 2", "Identity and profile", "OTP routes, profile state, compatibility mapping, vault design"],
            ["Phase 3", "Marketplace", "Explore flow, teaching services, mentor metadata, skill match logic"],
            ["Phase 4", "Exchange loop", "Session creation, room identifiers, wallet routes, transactions"],
            ["Phase 5", "Realtime and polish", "Socket signaling, dashboard views, charts, profile refinement"],
            ["Phase 6", "Hardening", "RLS, audit triggers, migration reconciliation, QA and documentation"],
        ],
        widths=[1.0, 1.6, 4.0],
    )
    add_detail_appendix(
        doc,
        "Appendix D: Testing Scenarios and Expected Observations",
        [
            "Happy-path registration",
            "Duplicate profile handling",
            "Admin-key fallback behavior",
            "Wallet insufficient-balance flow",
            "Mentor listing ownership checks",
            "Session-room creation verification",
            "RLS privacy boundary validation",
            "Realtime disconnect recovery",
            "Cross-origin configuration checks",
            "Post-deployment monitoring expectations",
        ],
        base_number="D",
    )

    path = OUT_DIR / "UNISKILL_SYNOPSIS_FINAL.docx"
    doc.save(path)
    return path


def build_srs(diagrams: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc, "UniSkill SRS")
    add_cover(
        doc,
        "Software Requirements Specification",
        "IEEE-style specification for the UniSkill platform",
        "SRS Report",
    )
    toc = [
        "1. Introduction",
        "2. Overall Description",
        "3. External Interface Requirements",
        "4. System Features and Functional Requirements",
        "5. Non-Functional Requirements",
        "6. Data Requirements and Data Dictionary",
        "7. Security, Privacy, and Compliance",
        "8. Validation, Acceptance, and Traceability",
        "9. Appendices",
    ]
    add_manual_toc(doc, toc)

    add_heading(doc, "1. Introduction", 1)
    sections = [
        ("1.1 Purpose", "This SRS defines the functional and non-functional requirements for UniSkill, a verified student skill exchange marketplace that supports profile verification, service listing, mentor discovery, booking, credits, and real-time collaboration."),
        ("1.2 Document Conventions", "Requirement identifiers use FR for functional requirements, NFR for non-functional requirements, and DR for data-related requirements. The words shall and must indicate mandatory behavior; should indicates a desirable but non-blocking behavior."),
        ("1.3 Intended Audience and Reading Suggestions", "This document is intended for project evaluators, developers, designers, testers, deployment engineers, and future maintainers. Readers may begin with overall description, then move to the system feature section and data dictionary."),
        ("1.4 Product Scope", "UniSkill enables a university student to create a trusted profile, declare skills and learning goals, publish or consume teaching services, request sessions, and exchange value through a controlled credits economy."),
    ]
    for head, text in sections:
        add_heading(doc, head, 2)
        add_para(doc, text)
        build_filler_block(doc, head, [
            ("clarity at the specification stage reduces ambiguity during implementation", "named entities, route flows, and persistent states", "teams can validate scope against observable behavior"),
            ("the platform spans interface and policy concerns", "frontend states, auth mechanics, and RLS-backed storage rules", "requirement quality must cover both experience and enforcement"),
        ])

    add_heading(doc, "2. Overall Description", 1)
    overall = {
        "2.1 Product Perspective": [
            "UniSkill is a web-native campus marketplace that integrates browser clients, SSR-aware server handlers, an Express backend, and Supabase-managed data services.",
            "The product may be deployed as a modular full-stack application, with the Next.js frontend and Express backend hosted independently while sharing identity, environment configuration, and origin policy agreements.",
        ],
        "2.2 Product Functions": [
            "Register student users and create baseline profiles.",
            "Verify identity through OTP and private vault storage.",
            "Expose public teaching signals and learning goals.",
            "Manage service listings, sessions, rooms, and credits.",
            "Maintain auditability, privacy boundaries, and future admin extensibility.",
        ],
        "2.3 User Classes and Characteristics": [
            "Learner: a student seeking help in an academic, technical, or extracurricular skill.",
            "Mentor: a student offering one or more services or teachable skills in exchange for credits.",
            "Admin or reviewer: a trusted governance role responsible for verification, moderation, and dispute review in future releases.",
        ],
        "2.4 Operating Environment": [
            "Modern browsers on desktop and mobile devices.",
            "Node.js runtime for frontend and backend services.",
            "Supabase-hosted PostgreSQL, auth, storage, and real-time features.",
        ],
        "2.5 Design and Implementation Constraints": [
            "The system depends on environment variables for auth keys and allowed origins.",
            "Sensitive verification data must not be exposed through public profile responses.",
            "The evolving schema requires compatibility-aware access logic during transition phases.",
        ],
        "2.6 User Documentation": [
            "Onboarding guide, verification instructions, profile setup help, mentor listing tips, and wallet usage notes should be available in-app or through linked documentation.",
        ],
        "2.7 Assumptions and Dependencies": [
            "Students possess valid email access and stable internet connectivity.",
            "The Supabase project remains available and correctly configured with migrations and RLS policies.",
            "Email delivery infrastructure for OTP or password reset messages remains operational.",
        ],
    }
    for head, bullets in overall.items():
        add_heading(doc, head, 2)
        add_bullets(doc, bullets)
        add_para(doc, "The implementation evidence in the repository supports these assumptions, but production hardening should make each dependency observable and testable in operations.")

    add_heading(doc, "3. External Interface Requirements", 1)
    add_heading(doc, "3.1 User Interfaces", 2)
    add_para(doc, "The system shall provide responsive interfaces for landing, signup, login, dashboard, explore, profile, wallet, and live room interactions. The visual language should remain clear enough for frequent campus use and flexible enough to distinguish learning mode from teaching mode.")
    add_table(
        doc,
        ["Interface", "Primary user goal", "Key inputs", "Expected outputs"],
        [
            ["Signup / Login", "Join and authenticate", "Email, password, OTP, identity details", "Authenticated session and initial profile state"],
            ["Profile", "Manage public presence", "Skills, goals, services, rate, bio", "Updated public data and mentor availability"],
            ["Explore", "Find mentors and offers", "Filters, searches, session selection", "Service and profile results"],
            ["Wallet", "Monitor or move credits", "Top-up and transfer requests", "Balance and transaction history"],
            ["Room", "Conduct live interaction", "Join room, messages, signaling", "Connected peer session"],
        ],
        widths=[1.3, 1.8, 2.1, 2.1],
    )
    add_heading(doc, "3.2 Hardware Interfaces", 2)
    add_para(doc, "UniSkill has no mandatory custom hardware interface. Optional camera, microphone, and speakers become relevant when live sessions use WebRTC-enabled room experiences.")
    add_heading(doc, "3.3 Software Interfaces", 2)
    add_bullets(doc, [
        "Supabase Auth for identity and browser/server session state",
        "Supabase PostgreSQL for tables, triggers, and policy-bound data access",
        "Socket.io for room events and lightweight live signaling",
        "Nodemailer or equivalent email dispatch service for OTP and reset flows",
        "Prisma schema for typed database modeling and reference consistency",
    ])
    add_heading(doc, "3.4 Communications Interfaces", 2)
    add_para(doc, "The system shall use HTTPS for standard application traffic and WebSocket-compatible transport for live room events. CORS configuration must permit only approved front-end origins, and session cookies must remain synchronized between SSR route handlers and the browser client.")
    add_figure(doc, diagrams["architecture"], "Figure 3.4: Cross-layer interface view connecting users, web clients, backend services, and Supabase", width=6.3)

    add_heading(doc, "4. System Features and Functional Requirements", 1)
    fr_groups = [
        ("4.1 FR1: User Registration and Login", [
            "FR1.1 The system shall allow a student to create an account using email and password credentials.",
            "FR1.2 The system shall create or reconcile a profile record associated with the authenticated user identifier.",
            "FR1.3 The system shall allow login through Supabase-authenticated credentials and return usable session information.",
            "FR1.4 The system shall provide a logout mechanism that clears active local session state.",
        ]),
        ("4.2 FR2: OTP and Identity Verification", [
            "FR2.1 The system shall send a six-digit OTP to an email or phone target for supported verification types.",
            "FR2.2 The system shall validate OTP value, type, and expiry before marking the verification step successful.",
            "FR2.3 The system shall expose verification status by user identifier for guided onboarding flows.",
            "FR2.4 The system shall maintain sensitive verification artifacts separately from public profile data.",
        ]),
        ("4.3 FR3: Profile Management", [
            "FR3.1 The system shall allow authenticated users to read and update their editable profile fields.",
            "FR3.2 The system shall support skills and learning goal arrays for discovery and matching.",
            "FR3.3 The system shall preserve compatibility across legacy and reconciled field names where required by the current codebase.",
            "FR3.4 The system shall prevent direct editing of protected fields such as credits through ordinary profile updates.",
        ]),
        ("4.4 FR4: Service Listing Management", [
            "FR4.1 The system shall allow a mentor to create, edit, read, and delete teaching service listings.",
            "FR4.2 A service listing shall contain title, category, description, duration, and price in credits.",
            "FR4.3 Creation of a teaching service shall mark the user as a mentor in profile state.",
        ]),
        ("4.5 FR5: Marketplace Discovery", [
            "FR5.1 The system shall expose mentor profiles for discovery in explore views.",
            "FR5.2 The system shall allow filtering by skill and minimum rating where supported.",
            "FR5.3 The system shall return service or mentor data ordered and shaped for interface consumption.",
        ]),
        ("4.6 FR6: Skill Match Recommendations", [
            "FR6.1 The system shall compute perfect matches when another user teaches what the current user wants and wants what the current user teaches.",
            "FR6.2 The system shall compute partial matches when another user teaches what the current user wants without a full reciprocal overlap.",
            "FR6.3 The system shall provide overlap details and rough match scores in response payloads.",
        ]),
        ("4.7 FR7: Session Booking", [
            "FR7.1 The system shall allow an authenticated learner to create a pending session with a selected mentor.",
            "FR7.2 The booking process shall generate a room identifier for subsequent room-based interaction.",
            "FR7.3 The system shall persist session status, title, scheduled time, duration, and price.",
        ]),
        ("4.8 FR8: Wallet and Credits", [
            "FR8.1 The system shall expose wallet balance and recent transactions for a user.",
            "FR8.2 The system shall support addition of credits through a server-side route that records the transaction metadata.",
            "FR8.3 The system shall support peer transfer of credits between two distinct users subject to balance checks.",
            "FR8.4 Transfer activity shall produce transaction records for both sender and receiver.",
        ]),
        ("4.9 FR9: Messaging and Live Room Signaling", [
            "FR9.1 The system shall allow clients to join a room and receive join and disconnect notifications.",
            "FR9.2 The system shall relay signal payloads for peer connection establishment.",
            "FR9.3 The system shall relay room-scoped chat messages with sender and timestamp metadata.",
            "FR9.4 The system shall broadcast explicit end-call events within a room.",
        ]),
        ("4.10 FR10: Reviews, Notifications, and Governance", [
            "FR10.1 The data model shall support reviews associated with sessions and reviewer-reviewee pairs.",
            "FR10.2 The data model shall support notifications with type, message, and read state.",
            "FR10.3 The platform shall remain extensible for admin review, moderation, and audit features.",
        ]),
    ]
    for heading, reqs in fr_groups:
        add_heading(doc, heading, 2)
        add_heading(doc, f"{heading}.1 Description and Priority", 3)
        add_para(doc, "Priority for this feature group is high because it contributes directly to the core exchange loop of trust, discovery, booking, session execution, or value transfer.")
        add_heading(doc, f"{heading}.2 Functional Requirements", 3)
        add_bullets(doc, reqs)
        add_heading(doc, f"{heading}.3 Rationale", 3)
        add_para(doc, "These requirements reflect observed repository behavior and the logical necessities of a campus learning marketplace. Together they define the minimum coherent operating surface for UniSkill.")

    add_figure(doc, diagrams["sequence"], "Figure 4.10: Booking and credit reservation sequence used to derive session-related requirements", width=6.5)

    add_heading(doc, "5. Non-Functional Requirements", 1)
    nfrs = [
        ("Performance", [
            "NFR1 The marketplace and dashboard should render primary data within acceptable interactive latency under normal campus loads.",
            "NFR2 OTP verification responses should complete fast enough to avoid user abandonment during signup.",
            "NFR3 Room signaling events should propagate with low delay suitable for live session setup.",
        ]),
        ("Security", [
            "NFR4 Sensitive verification artifacts must remain inaccessible to unauthorized users.",
            "NFR5 Profile, session, transaction, and message access shall respect row-level or route-level authorization checks.",
            "NFR6 Secrets, service role keys, and email credentials must not be exposed to browser clients.",
        ]),
        ("Reliability", [
            "NFR7 Critical writes such as session creation and wallet transfer should avoid partial failure states where possible.",
            "NFR8 The system should degrade gracefully if email, socket, or one auxiliary service is temporarily unavailable.",
        ]),
        ("Usability", [
            "NFR9 The product shall remain understandable for first-time student users without formal training.",
            "NFR10 The UI should preserve readable layouts across mobile and desktop widths.",
        ]),
        ("Scalability", [
            "NFR11 The architecture should support modular growth in users, listings, and session traffic without requiring a full rewrite.",
            "NFR12 Data indexes, policy design, and API partitioning should anticipate campus-scale load expansion.",
        ]),
        ("Maintainability", [
            "NFR13 Clear route boundaries and schema migration history shall support future maintenance.",
            "NFR14 Compatibility handling should be documented until legacy field drift is eliminated.",
        ]),
    ]
    for head, items in nfrs:
        add_heading(doc, f"5.{nfrs.index((head, items)) + 1} {head}", 2)
        add_bullets(doc, items)
        add_para(doc, f"{head} requirements are essential because UniSkill is not only a demo interface; it coordinates identity, discovery, and value movement for real users.")

    add_heading(doc, "6. Data Requirements and Data Dictionary", 1)
    add_heading(doc, "6.1 Core Data Entities", 2)
    add_figure(doc, diagrams["erd"], "Figure 6.1: Core persistent entities in the UniSkill domain", width=6.6)
    add_heading(doc, "6.2 Data Dictionary", 2)
    rows = [
        ["profiles", "id", "UUID", "Primary user identifier linked to auth user"],
        ["profiles", "skills / learning_goals", "TEXT[]", "Public learning and teaching metadata"],
        ["user_verification_vault", "college_email", "TEXT", "Private verified institution contact"],
        ["sessions", "status", "TEXT", "Lifecycle state for a booked engagement"],
        ["sessions", "room_id", "TEXT", "Logical room key for live communication"],
        ["transactions", "amount", "INTEGER", "Credit movement value"],
        ["service_listings", "price_credits", "INTEGER", "Credits required for a mentor service"],
        ["messages", "content", "TEXT", "Direct communication payload"],
        ["notifications", "data", "JSONB", "Auxiliary structured metadata"],
    ]
    add_table(doc, ["Table", "Field", "Type", "Description"], rows, widths=[1.4, 1.8, 1.0, 3.0])
    add_heading(doc, "6.3 Data Integrity Rules", 2)
    add_numbered(doc, [
        "Session rows shall reference valid mentor and student profile identifiers.",
        "Transaction rows for a peer transfer should be paired logically for sender and receiver visibility.",
        "Verification vault entries shall reference the same user identifier as their public profile counterpart.",
        "Service listings shall only be editable by their owning mentor.",
    ])

    add_heading(doc, "7. Security, Privacy, and Compliance", 1)
    add_para(doc, "Security requirements in UniSkill are unusually central because the platform blends private identity proofing with public mentor discovery. The schema and route structure therefore distinguish discoverability from confidentiality.")
    add_table(
        doc,
        ["Concern", "Current design approach", "Required assurance"],
        [
            ["Profile privacy", "Public profile table plus private verification vault", "Sensitive identity artifacts never leak through public browse flows"],
            ["Session authorization", "User-bound session policies and authenticated access", "Only session participants can inspect protected session data"],
            ["Wallet control", "Server-side balance checks and transaction logging", "No arbitrary client-side balance mutation"],
            ["Auditability", "Audit log trigger on profile changes", "Administrative review trail exists for sensitive mutations"],
        ],
        widths=[1.5, 2.5, 2.4],
    )
    add_heading(doc, "7.1 Compliance Considerations", 2)
    add_bullets(doc, [
        "Minimize exposure of personally identifiable student data.",
        "Keep clear consent boundaries for email, phone, and identity document use.",
        "Support institutional review through logs and moderation paths where needed.",
    ])

    add_heading(doc, "8. Validation, Acceptance, and Traceability", 1)
    add_heading(doc, "8.1 Validation Strategy", 2)
    add_bullets(doc, [
        "Requirement walkthroughs against implemented routes and pages",
        "Scenario testing for signup, verification, mentor publishing, booking, and transfer",
        "Regression checks whenever schema reconciliation changes compatibility behavior",
        "Security-focused checks for private table exposure and policy bypass risk",
    ])
    add_heading(doc, "8.2 Acceptance Criteria", 2)
    add_numbered(doc, [
        "A student can create an account and retrieve an authenticated profile state.",
        "A student can publish teaching metadata and see discoverable mentor-facing results.",
        "A learner can create a pending session tied to a room identifier.",
        "A user can inspect wallet data and complete a valid credit transfer.",
        "Realtime room events can be exchanged between participants.",
    ])
    add_heading(doc, "8.3 Requirement Traceability Summary", 2)
    add_table(
        doc,
        ["Requirement group", "Primary source artifact", "Verification method"],
        [
            ["FR1-FR3", "Auth routes, OTP routes, AuthContext", "Integration and UI flow tests"],
            ["FR4-FR7", "Teaching-services route, sessions route, profile and explore pages", "Scenario tests and DB inspection"],
            ["FR8", "Wallet routes and dashboard/wallet UI", "Balance and transaction assertions"],
            ["FR9", "Socket.io setup and room page", "Live room event verification"],
            ["Security / Data", "Supabase migrations and RLS policies", "Policy review and guarded access tests"],
        ],
        widths=[1.6, 2.6, 2.2],
    )

    add_heading(doc, "9. Appendices", 1)
    for idx, title in enumerate([
        "Glossary of domain terms",
        "Environment configuration summary",
        "Assumed production deployment topology",
        "Known schema compatibility considerations",
        "Potential future requirement additions",
    ], start=1):
        add_heading(doc, f"9.{idx} {title}", 2)
        build_filler_block(doc, title, [
            ("appendices provide durable handoff context for future maintainers", "explicit assumptions and cross-layer notes", "teams avoid rediscovering the same implementation constraints"),
            ("UniSkill is still evolving in a realistic engineering environment", "migration history, compatibility mapping, and staged features", "supplementary notes remain operationally valuable"),
        ])

    add_heading(doc, "Appendix A: Detailed Use Case Specifications", 1)
    use_cases = [
        ("A1", "Student Registration", "Prospective student", ["No active account for the same identity"], ["Open signup page", "Enter personal details", "Submit registration request", "Receive authenticated account or fallback guidance", "Proceed to verification steps"], ["Baseline profile exists and user can continue onboarding"]),
        ("A2", "OTP Verification", "Authenticated student", ["User has a valid target email or phone", "OTP request has been generated"], ["Request verification code", "Receive OTP", "Enter six-digit code", "System validates type and expiry", "Verification step advances"], ["Verification status is updated"]),
        ("A3", "Publish Mentor Service", "Mentor", ["User is authenticated", "Profile contains at least one teachable skill"], ["Open profile page", "Fill service title, category, duration, and price", "Submit service form", "System validates values", "Service becomes available in listings"], ["Service listing is stored and mentor flag remains true"]),
        ("A4", "Browse and Select Mentor", "Learner", ["Learner is authenticated"], ["Open explore page", "Search or filter mentor results", "Inspect profile or service details", "Compare rates and focus areas", "Choose a mentor or service"], ["Learner has enough information to initiate booking"]),
        ("A5", "Book Session", "Learner", ["Learner is authenticated", "Target mentor exists"], ["Submit booking data", "Server creates pending session", "Room identifier is attached", "Client receives success response", "Learner sees pending status"], ["Session row is persisted"]),
        ("A6", "Transfer Credits", "Student", ["Sender and receiver are distinct users", "Sender balance is sufficient"], ["Submit transfer request", "Server validates users and amount", "Sender balance decreases", "Receiver balance and earnings increase", "Transactions are logged"], ["Economic state is updated for both parties"]),
        ("A7", "Join Live Room", "Session participant", ["Session exists with room identifier"], ["Connect to socket server", "Join room", "Exchange signal payloads", "Send chat messages", "Optionally end call"], ["Participants share a synchronized live room state"]),
        ("A8", "Review Session", "Student", ["Completed session exists"], ["Open review form", "Choose rating", "Submit comment", "Persist review record"], ["Mentor or learner receives quality feedback"]),
    ]
    for args in use_cases:
        add_use_case_spec(doc, *args)

    add_heading(doc, "Appendix B: Expanded Requirement Traceability Matrix", 1)
    matrix_rows = []
    for idx in range(1, 21):
        matrix_rows.append([
            f"REQ-{idx:02d}",
            ["Auth", "Verification", "Profile", "Services", "Marketplace", "Sessions", "Wallet", "Realtime", "Security", "Data"][idx % 10],
            "Repository route, page, migration, or schema artifact",
            "Integration / UI / policy verification",
        ])
    add_table(doc, ["Requirement ID", "Domain", "Primary artifact source", "Validation mode"], matrix_rows, widths=[1.0, 1.3, 2.7, 1.8])

    add_detail_appendix(
        doc,
        "Appendix C: Error Handling, Boundary Conditions, and Recovery Requirements",
        [
            "Invalid credential handling",
            "Email-not-confirmed responses",
            "OTP expiry and retry limitations",
            "Profile-auth desynchronization recovery",
            "Missing service-listing ownership",
            "Wallet insufficient funds",
            "Network interruption during room join",
            "Unauthorized route access",
            "Invalid environment configuration",
            "Migration mismatch and compatibility fallback",
        ],
        base_number="C",
    )
    add_detail_appendix(
        doc,
        "Appendix D: Detailed Data Rules and Quality Constraints",
        [
            "Profile field semantics",
            "Verification vault semantics",
            "Session status semantics",
            "Transaction classification rules",
            "Review integrity rules",
            "Notification state semantics",
            "Service listing validation rules",
            "Message ownership rules",
            "Favorites consistency rules",
            "Audit logging expectations",
        ],
        base_number="D",
    )

    path = OUT_DIR / "UNISKILL_SRS_FINAL.docx"
    doc.save(path)
    return path


def build_uml(diagrams: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc, "UniSkill UML and Design")
    add_cover(
        doc,
        "UML and Design Documentation",
        "Analysis model, design model, diagrams, dictionaries, and screen-oriented architecture notes for UniSkill",
        "UML / Design Report",
    )
    toc = [
        "1. Introduction",
        "2. Analysis Model",
        "3. Design Model",
        "4. Database Design",
        "5. Component and Interface Design",
        "6. Dynamic Behavior Models",
        "7. Screen and Interaction Notes",
        "8. Design Decisions and Future Evolution",
    ]
    add_manual_toc(doc, toc)

    add_heading(doc, "1. Introduction", 1)
    for head, text in [
        ("1.1 Purpose", "This document presents the analysis and design model of UniSkill using UML-style views, data flow abstractions, and architecture-level diagrams derived from the implemented repository and its intended product direction."),
        ("1.2 Document Conventions", "Diagrams are presented as logical views rather than exact code-generation artifacts. Entity names mirror repository terminology wherever possible."),
        ("1.3 Intended Audience and Reading Suggestions", "Developers, reviewers, database designers, and evaluators should use this document to understand how UniSkill behaves structurally, dynamically, and visually."),
        ("1.4 References", "References include the project repository, route handlers, database migration files, and official framework documentation used by the current implementation."),
    ]:
        add_heading(doc, head, 2)
        add_para(doc, text)

    add_heading(doc, "2. Analysis Model", 1)
    add_heading(doc, "2.1 Methodology Used", 2)
    add_para(doc, "The analysis model combines user-centric use cases, data-oriented conceptual modeling, and control-flow reasoning. This blend suits UniSkill because the platform must explain not only what the user sees, but how trust, exchange, and persistence interact under the surface.")
    add_heading(doc, "2.2 Use Case Diagram", 2)
    add_figure(doc, diagrams["use_case"], "Figure 2.2: Major actors and use cases in the UniSkill platform", width=6.5)
    add_para(doc, "The use case view emphasizes that students act in multiple roles. A single user may learn, teach, publish services, move credits, and join rooms, while administrative review remains a secondary but important trust function.")
    add_heading(doc, "2.3 ER Model", 2)
    add_figure(doc, diagrams["erd"], "Figure 2.3: ER model for profiles, sessions, transactions, messages, and service listings", width=6.5)
    add_heading(doc, "2.4 Data Flow Diagram", 2)
    add_figure(doc, diagrams["dfd"], "Figure 2.4: Data flow between student actors, processing units, and data stores", width=6.3)
    add_heading(doc, "2.5 Control Flow Diagram", 2)
    add_figure(doc, diagrams["activity"], "Figure 2.5: Mentor onboarding control flow", width=6.2)
    add_heading(doc, "2.6 State Transition Diagram", 2)
    add_figure(doc, diagrams["state"], "Figure 2.6: Session state progression and exception branches", width=6.2)
    add_heading(doc, "2.7 Narrative Analysis", 2)
    build_filler_block(doc, "analysis modeling", [
        ("UniSkill depends on transitions between trusted states, not on static profile pages alone", "verification steps, mentor publication, booking state, and credit movement", "diagrammatic analysis is a necessary design aid"),
        ("the system has both event-driven and data-driven behavior", "socket events, REST-style writes, and policy-protected reads", "a single UML notation would not explain the product adequately by itself"),
        ("students move fluidly between platform roles", "shared profile identity with contextual behavior", "actor analysis must account for dual-role journeys"),
    ])

    add_heading(doc, "3. Design Model", 1)
    add_heading(doc, "3.1 Architectural Design", 2)
    add_heading(doc, "3.1.1 System Architecture Diagram", 3)
    add_figure(doc, diagrams["architecture"], "Figure 3.1.1: System architecture diagram for UniSkill", width=6.5)
    add_heading(doc, "3.1.2 Description of Architectural Design", 3)
    for text in [
        "The architecture separates presentation, route mediation, business logic, realtime event transport, and persistence. This keeps the student-facing experience responsive while allowing secure server-side coordination of stateful operations.",
        "Next.js pages and SSR-aware route handlers own browser-facing behavior. Express encapsulates server-level services such as OTP, wallet logic, and cross-origin socket configuration. Supabase anchors storage, identity, and policy enforcement.",
    ]:
        add_para(doc, text)

    add_heading(doc, "3.2 Database Design", 2)
    add_heading(doc, "3.2.1 Data Dictionary", 3)
    add_table(
        doc,
        ["Entity", "Responsibility", "Key relationships"],
        [
            ["profiles", "Public user identity and participation state", "Related to sessions, transactions, reviews, messages, favorites"],
            ["user_verification_vault", "Private identity proofing and institutional contact data", "One-to-one with profile/auth user"],
            ["sessions", "Booked learning engagements", "Connects learner, mentor, room, reviews, and transactions"],
            ["transactions", "Credit movement log", "Belongs to user and optionally a session"],
            ["service_listings", "Bookable mentor offers", "Belongs to mentor profile"],
        ],
        widths=[1.4, 2.5, 2.4],
    )
    add_heading(doc, "3.2.2 Normalization", 3)
    build_filler_block(doc, "normalization", [
        ("core transactional data is mostly decomposed into separate entities with explicit references", "sessions, reviews, transactions, and messages", "write boundaries remain understandable"),
        ("some compatibility fields intentionally duplicate business meaning during migration periods", "name versus display_name and credits versus credits_balance", "normalization discussion should acknowledge pragmatic transitional design"),
    ])

    add_heading(doc, "3.3 Component Design", 2)
    add_heading(doc, "3.3.1 Flow Chart", 3)
    add_figure(doc, diagrams["component"], "Figure 3.3.1: Component relationship flow for primary UniSkill modules", width=6.4)
    add_heading(doc, "3.3.2 Component Responsibilities", 3)
    add_numbered(doc, [
        "AuthContext bridges authenticated identity, compatibility mapping, and verification-state hydration for the client.",
        "Teaching-service routes encapsulate mentor-owned CRUD logic and profile mentor-state updates.",
        "Session routes translate booking intent into persistent session records and room identifiers.",
        "Wallet routes own balance lookup, credit addition, and transfer record generation.",
        "Socket.io handlers coordinate join, signal, chat, disconnect, and call-end room behavior.",
    ])

    add_heading(doc, "3.4 Interface Design", 2)
    add_heading(doc, "3.4.1 Screenshots and Surface Mapping", 3)
    build_filler_block(doc, "screen design", [
        ("the landing page positions UniSkill as a verified campus network", "hero statistics, feature cards, and action buttons", "the UI communicates trust and momentum before the user signs up"),
        ("dashboard and home overview pages summarize session rhythm and credit movement", "activity charts, quick links, and recent sessions", "design aligns with repeat-use student behavior"),
        ("profile and explore pages are task-oriented rather than merely descriptive", "service authoring forms, discoverable filters, and public mentor signals", "interface design is tied directly to transaction readiness"),
    ])

    add_heading(doc, "4. Database Design", 1)
    add_heading(doc, "4.1 Table Relationships", 2)
    add_para(doc, "The relationship graph is anchored by profiles. Sessions establish bidirectional participation between student and mentor. Transactions create economic traceability, while reviews and messages create qualitative and conversational continuity.")
    add_heading(doc, "4.2 Security Policy Notes", 2)
    add_bullets(doc, [
        "Profiles are publicly selectable for discovery-oriented use cases.",
        "Verification vault rows are private to the owning authenticated user.",
        "Transactions are viewable by the owning user, with broader control reserved for service-role contexts.",
        "Messages and notifications remain user-scoped.",
    ])
    add_heading(doc, "4.3 Trigger and Automation Design", 2)
    build_filler_block(doc, "triggers and automation", [
        ("signup automation reduces profile bootstrap friction", "on_auth_user_created and profile initialization", "the design links auth creation to domain readiness"),
        ("audit triggers increase accountability for profile mutation", "user_audit_log inserts", "design governance becomes concrete rather than aspirational"),
        ("updated_at triggers standardize freshness semantics across mutable tables", "profile and session maintenance", "timestamp usage stays reliable"),
    ])

    add_heading(doc, "5. Component and Interface Design", 1)
    add_heading(doc, "5.1 Realtime Collaboration View", 2)
    add_figure(doc, diagrams["sequence"], "Figure 5.1: Sequence view for booking and follow-up room interaction", width=6.5)
    add_heading(doc, "5.2 API Surface Summary", 2)
    add_table(
        doc,
        ["Endpoint family", "Representative path", "Purpose"],
        [
            ["Auth", "/api/auth/register", "Create user and reconcile profile state"],
            ["OTP", "/api/otp/send", "Dispatch verification code"],
            ["Wallet", "/api/wallet/transfer", "Move credits between users"],
            ["Skills", "/api/skills/match/:userId", "Compute reciprocal learning opportunities"],
            ["Services", "/api/teaching-services", "Manage mentor service offers"],
            ["Sessions", "/api/sessions", "Create pending bookings"],
        ],
        widths=[1.2, 2.0, 3.0],
    )
    add_heading(doc, "5.3 UI Surface Design Notes", 2)
    for note in [
        "Explore mode and freelance mode indicate that the product supports distinct but related user intentions.",
        "Profile management intentionally brings skill curation and service authoring into the same operational surface.",
        "The room page acts as a session execution boundary where previously independent booking and communication flows converge.",
    ]:
        add_para(doc, note)

    add_heading(doc, "6. Dynamic Behavior Models", 1)
    add_heading(doc, "6.1 Booking Sequence", 2)
    add_para(doc, "Booking begins when a learner selects a mentor or service and ends when a pending session row and room identifier are returned to the interface. Downstream wallet reservation and final settlement logic can deepen this sequence in future revisions.")
    add_heading(doc, "6.2 Mentor Onboarding Activity", 2)
    add_figure(doc, diagrams["activity"], "Figure 6.2: Mentor onboarding activity view", width=5.8)
    add_heading(doc, "6.3 Session State Machine", 2)
    add_figure(doc, diagrams["state"], "Figure 6.3: Session state machine", width=6.0)
    add_heading(doc, "6.4 Event Model Narrative", 2)
    build_filler_block(doc, "event behavior", [
        ("user-visible progress depends on a sequence of explicit state transitions", "pending sessions, room joins, and transfer records", "dynamic models are necessary for test design"),
        ("realtime behavior supplements but does not replace persistent data", "socket events and stored rows", "the system remains recoverable after transient disconnects"),
    ])

    add_heading(doc, "7. Screen and Interaction Notes", 1)
    screen_sections = [
        "Landing and value proposition communication",
        "Signup and login interaction structure",
        "Profile authoring and verification cues",
        "Explore filters, cards, and conversion actions",
        "Wallet transparency and transaction visibility",
        "Room interaction with peer connection events",
    ]
    for idx, title in enumerate(screen_sections, start=1):
        add_heading(doc, f"7.{idx} {title}", 2)
        build_filler_block(doc, title, [
            ("interface choices shape trust as much as backend policies do", "surface hierarchy, readable metrics, and visible state", "screen notes belong in technical documentation"),
            ("UniSkill is intended for repeated campus use rather than occasional novelty visits", "quick actions, summary cards, and operational forms", "interaction design must stay efficient"),
        ])

    add_heading(doc, "8. Design Decisions and Future Evolution", 1)
    add_heading(doc, "8.1 Notable Current Design Decisions", 2)
    add_numbered(doc, [
        "Use Supabase for identity, storage, and RLS rather than building custom auth and policy infrastructure from scratch.",
        "Preserve a compatibility layer while the schema transitions from early and reconciled field sets.",
        "Split public profile data from private verification data to protect sensitive information.",
        "Keep realtime behavior lightweight and room-scoped, with persistence handled separately by database writes.",
    ])
    add_heading(doc, "8.2 Future Evolution Paths", 2)
    add_bullets(doc, [
        "Richer admin moderation dashboards and dispute flows",
        "Full session settlement logic with stronger transactional guarantees",
        "Recommendation and ranking models for mentor discovery",
        "Analytics surfaces for learner progress and campus engagement",
        "Cross-campus tenancy and institution-specific verification rules",
    ])
    add_heading(doc, "8.3 Closing Design Summary", 2)
    add_para(doc, "The UniSkill design is strongest when read as a coherent service ecosystem rather than as disconnected pages or tables. Its diagrams show a project that already has real structure and clear room for disciplined growth.")

    add_detail_appendix(
        doc,
        "Appendix A: Detailed Entity and Relationship Commentary",
        [
            "Profile identity decomposition",
            "Verification vault separation rationale",
            "Session ownership and lifecycle commentary",
            "Transaction pairing logic",
            "Review authorship and reputation impact",
            "Message privacy and read-state design",
            "Notification payload design",
            "Service-listing evolution path",
            "Audit-log governance role",
            "Compatibility-field retirement strategy",
        ],
        base_number="A",
    )
    add_heading(doc, "Appendix B: Use Case Realization Notes", 1)
    for args in [
        ("B1", "Registration Realization", "Student", ["Signup surface is available"], ["Frontend captures values", "Auth route attempts admin create or signup fallback", "Profile is read or created", "AuthContext hydrates client state"], ["User reaches a coherent starting profile"]),
        ("B2", "Verification Realization", "Student", ["User has a target for OTP"], ["OTP route stores code", "Email service dispatches message", "Verify route checks code", "Status endpoint reflects progress"], ["Trust state advances"]),
        ("B3", "Booking Realization", "Learner", ["Mentor and service context exist"], ["Explore flow exposes choices", "Session API inserts row", "Room identifier is attached", "Dashboard can later surface recent activity"], ["Booking artifact exists"]),
        ("B4", "Transfer Realization", "Student", ["Sender balance is adequate"], ["Wallet route reads sender", "Wallet route reads receiver", "Balances mutate", "Transaction rows are inserted"], ["Economic narrative is persisted"]),
        ("B5", "Room Realization", "Participant", ["Room identifier exists"], ["Socket join-room fires", "Signal and chat events propagate", "Disconnect or end-call notifications broadcast"], ["Interaction state remains synchronized"]),
    ]:
        add_use_case_spec(doc, *args)
    add_detail_appendix(
        doc,
        "Appendix C: Interface-State and Screenflow Commentary",
        [
            "Landing to signup transition",
            "Signup to verification transition",
            "Verification to profile-completion transition",
            "Profile to explore transition",
            "Explore to booking transition",
            "Booking to dashboard transition",
            "Dashboard to wallet transition",
            "Dashboard to room transition",
            "Profile to mentor-publishing transition",
            "Session closure to review transition",
        ],
        base_number="C",
    )
    add_detail_appendix(
        doc,
        "Appendix D: Design Alternatives and Trade-off Discussion",
        [
            "Centralized backend versus serverless-heavy design",
            "Public profile breadth versus privacy minimization",
            "Wallet credits versus direct payments",
            "Service listings versus skill-only discovery",
            "RLS-first data protection versus app-only authorization",
            "Incremental schema compatibility versus forced migration cleanup",
            "Socket-based signaling versus polling",
            "Campus-only identity versus open marketplace enrollment",
            "Managed services versus self-hosted infrastructure",
            "Feature-rich room experience versus minimal coordination surface",
        ],
        base_number="D",
    )

    path = OUT_DIR / "UNISKILL_UML_FINAL.docx"
    doc.save(path)
    return path


def main() -> None:
    ensure_dir(OUT_DIR)
    ensure_dir(ASSET_DIR)
    diagrams = make_diagrams()
    paths = [
        build_synopsis(diagrams),
        build_srs(diagrams),
        build_uml(diagrams),
    ]
    print("Generated:")
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
